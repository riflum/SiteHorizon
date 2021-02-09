# PyQt5
from PyQt5.QtWidgets import (QMessageBox, QApplication, QComboBox, QPushButton, 
                                QMainWindow, QLabel, QDialog, QGridLayout, QVBoxLayout, 
                                QTableView, QLineEdit, QFileDialog)
from PyQt5.QtCore import QFile, Qt, QSortFilterProxyModel, QRegularExpression
from PyQt5 import uic
from PyQt5.QtSql import QSqlQuery, QSqlTableModel, QSqlRelationalTableModel, QSqlRelation
# python library
import sys
from itertools import chain
import logging
# site and horizon module
from SiteWindow import SiteWindow as SW
import SiteWindow
from HorizonWindow import HorizonWindow
import connection
# from GetDataFromHorizonAll import getData
# python-docx
from docx import Document
from docx.shared import Cm, Pt, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Main UI file
mainFile = "UI\main.ui"
Ui_MainWindow, QtBaseClass = uic.loadUiType(mainFile)

### Logging - set the file to log and basic level
# format, datefmt, filename, level, filemode 
logging.basicConfig(format='%(asctime)s - %(levelname)s: %(message)s',datefmt='%Y/%m/%d %I:%M:%S %p',
                    level=logging.DEBUG, filemode="w")
# logging.basicConfig(level=logging.DEBUG)
logging.info("Warning")

def parseRange(rng):
    """ this method is used for parse range of row/page using dash "-" """
    parts = rng.split("-")
    if 1 > len(parts) > 2:
        raise ValueError ("Bad range: '%s'" %s (rng, ))
    parts = [int (i) for i in parts]
    start = parts[0]
    end = start if len(parts) == 1 else parts[1]
    if start > end:
        end, start = start, end
    return range(start, end + 1)

def parseRangeList(rngs):
    """ This method is for parse range of row/page and insertit as list"""
    return sorted(set(chain(*[parseRange(rng) for rng in rngs.split(',')])))
   
def initializeModel(model, table):
    """ This is use to initialize model with select database table """

    model.setTable(table)
    model.setEditStrategy(QSqlTableModel.OnManualSubmit)
    model.select()

def initializeTableEditModel(model, table):
    model.setTable(table)
    model.setEditStrategy(QSqlTableModel.OnFieldChange)
    model.select()

def proxyModelRowIndex(model, text, textLanguage=2):
    """ This method function is for getting the result of match function from site database tables
        to another related database tables """
    
    # initialize proxy model
    proxy = QSortFilterProxyModel()
    # set the proxy source model
    proxy.setSourceModel(model)
    # fetch more data from model/database so we can get data more than 255 row
    if (proxy.canFetchMore(model.index(model.rowCount(), model.columnCount()))):
        proxy.fetchMore(model.index(model.rowCount(), model.columnCount()))
    proxy.setFilterCaseSensitivity(Qt.CaseSensitive)
    # Filter the desired text/value
    proxy.setFilterFixedString(text)
    # matching the value with source model, start at index(0,0)
    matchingIndex = proxy.mapToSource(proxy.index(0,0))
    # validation if the data is not empty or "" get the row of the data
    # if not return index (-1) or empty
    if text != "":
        rowKindObs = matchingIndex.row()
    else:
        rowKindObs = -1
    # rowKindObs = matchingIndex.row()
    # get the data from the source model based on row and column
    result = model.data(model.index(rowKindObs, textLanguage))
    return result

def proxyModelRowIndexRegXMunsell(model, text, textLanguage=2):
    """ This method function is for getting the result of match function from site database tables
        to another related database tables """
    
    # initialize proxy model
    proxy = QSortFilterProxyModel()
    # set the proxy source model
    proxy.setSourceModel(model)
    # fetch more data from model/database so we can get data more than 255 row
    if (proxy.canFetchMore(model.index(model.rowCount(), model.columnCount()))):
        proxy.fetchMore(model.index(model.rowCount(), model.columnCount()))
    proxy.setFilterCaseSensitivity(Qt.CaseSensitive)
    # Filter Column
    proxy.setFilterKeyColumn(4)
    # Filter the desired text/value
    re = QRegularExpression("^{}$".format(text))
    re1 = QRegularExpression("^%s$" % (text))
    
    proxy.setFilterRegularExpression(re1)
    # matching the value with source model, start at index(0,0)
    matchingIndex = proxy.mapToSource(proxy.index(0,0))
    # validation if the data is not empty or "" get the row of the data
    # if not return index (-1) or empty
    if text != "":
        rowKindObs = matchingIndex.row()
    else:
        rowKindObs = -1
    # rowKindObs = matchingIndex.row()
    # get the data from the source model based on row and column
    result = model.data(model.index(rowKindObs, textLanguage))
    return result
def proxyModelRowIndexRegX(model, text, textLanguage=2):
    """ This method function is for getting the result of match function from site database tables
        to another related database tables """
    
    # initialize proxy model
    proxy = QSortFilterProxyModel()
    # set the proxy source model
    proxy.setSourceModel(model)
    # fetch more data from model/database so we can get data more than 255 row
    if (proxy.canFetchMore(model.index(model.rowCount(), model.columnCount()))):
        proxy.fetchMore(model.index(model.rowCount(), model.columnCount()))
    proxy.setFilterCaseSensitivity(Qt.CaseSensitive)
    # Filter the desired text/value
    re = QRegularExpression("^{}$".format(text))
    re1 = QRegularExpression("^%s$" % (text))
    
    proxy.setFilterRegularExpression(re1)
    # matching the value with source model, start at index(0,0)
    matchingIndex = proxy.mapToSource(proxy.index(0,0))
    # validation if the data is not empty or "" get the row of the data
    # if not return index (-1) or empty
    if text != "":
        rowKindObs = matchingIndex.row()
    else:
        rowKindObs = -1
    # rowKindObs = matchingIndex.row()
    # get the data from the source model based on row and column
    result = model.data(model.index(rowKindObs, textLanguage))
    return result

class First(QMainWindow, Ui_MainWindow):
    def __init__(self, parent=None):
        super(First, self).__init__(parent)
        # setup the "main.ui"
        self.setupUi(self)
        # TODO: Try to change 2 or more model that access the samae database table
        # to one of model and then the other model change to proxymodel.
        self.siteModel = QSqlTableModel()
        self.horizonModel = QSqlTableModel()
        # initilize site model
        initializeModel(self.siteModel, "Site")
        # initilize horizon model
        initializeModel(self.horizonModel, "HorizonAll")
        # populate the model required
        self.populateModel()

        ### Model for editiong Table of Site or Horizon
        self.siteTableModel = QSqlTableModel()
        self.horizonTableModel = QSqlTableModel()
        # initialize site and horizon table edit
        initializeTableEditModel(self.siteTableModel, "Site")
        initializeTableEditModel(self.horizonTableModel, "HorizonAll")
        

        # setting the SiteWindow as a dialog 
        self.siteDialog = SW()
        # setting the window modality to Qt.ApplicationModal so the main window can't 
        # force close the main window if the Site Window active/show
        self.siteDialog.setWindowModality(Qt.ApplicationModal)
        # setting the HorizonWindow to a dialog
        self.horizonDialog = HorizonWindow()
        # setting the window modality to Qt.ApplicationModal so the main window can't 
        # force close the main window if the Horizon Window active/show
        self.horizonDialog.setWindowModality(Qt.ApplicationModal)

        ### ACTION
        # Open Site Entry
        self.Action_Site.triggered.connect(self.siteClicked)
        # Open Site Entry
        self.Action_Horizon.triggered.connect(self.horizonClicked)
        # Close the application
        self.Action_Quit.triggered.connect(self.close)
        # export to docx event triggered
        self.Action_Export_to_Docx.triggered.connect(self.exportToDocx)
        # backup database
        self.Action_Backup_Database.triggered.connect(self.backupDatabaseDialog)
        # delete site and database data
        self.Action_Delete.triggered.connect(self.deleteTablesData)
        # SHPy version
        self.Action_About.triggered.connect(self.aboutSHPy)

        ### EVENT
        # open the site window for entry data
        self.Site_PB.clicked.connect(self.siteClicked)
        # open the site window for entry data
        self.Horizon_PB.clicked.connect(self.horizonClicked)
        # just for check something - "delete later!"
        self.Check_PB.clicked.connect(self.reloadClicked)
        # expot to docx pushButton
        self.Export_To_Docx_PB.clicked.connect(self.exportToDocx)
        # button to show the table
        self.Table_PB.clicked.connect(self.showTable)
        
        # # I thik it doesnt work if the data in table (foreign key) is empty/not relatable with 
        # # The result is no Table showed
        # self.initRelationalTable()

    def backupDatabaseDialog(self):
        
        logging.info("backup dialog")

        self.backupDatabaseDialog = QDialog()
        gridLayout = QGridLayout()
        newFileLabel = QLabel("Database Name")
        self.newFileLineEdit = QLineEdit("SiteHorizon.db")
        self.newFileLineEdit.setReadOnly(True)
        self.saveSqliteDBLineEdit = QLineEdit()
        self.saveSqliteDBLineEdit.setReadOnly(True)
        saveSqliteDBButton = QPushButton("Database Save Location")
        okPushButton = QPushButton("OK")
        cancelPushButton = QPushButton("Cancel")

        gridLayout.setColumnStretch(1, 1)
        # set the minimum width of '0' / first grid column
        gridLayout.setColumnMinimumWidth(0, 120)

        gridLayout.addWidget(newFileLabel, 0, 0)
        gridLayout.addWidget(self.newFileLineEdit, 0, 1)
        gridLayout.addWidget(self.saveSqliteDBLineEdit, 1, 0, 1, 2)
        gridLayout.addWidget(saveSqliteDBButton, 2, 0, 1, 2)
        gridLayout.addWidget(okPushButton, 3, 0)
        gridLayout.addWidget(cancelPushButton, 3, 1)

        okPushButton.clicked.connect(self.okBackupDatabase)
        cancelPushButton.clicked.connect(self.backupDatabaseDialog.close)
        saveSqliteDBButton.clicked.connect(self.setSaveLocation)

        self.backupDatabaseDialog.setLayout(gridLayout)
        self.backupDatabaseDialog.setWindowTitle("Backup Database")
        self.backupDatabaseDialog.exec()

    def okBackupDatabase(self):
        logging.info("Ok to backup")

        ### Backup method
        # Begin imeddiate method to lock the database for access
        query = QSqlQuery()
        query.prepare("BEGIN IMMEDIATE;")
        query.exec()

        file = QFile()
        # Copy/backup database from database name to new save database locations
        file.copy("Database\SiteHorizon.db", self.saveSqliteDBLineEdit.text())
        
        # roleback/ unlcok the access
        query.prepare("ROLLBACK;")
        query.exec()

        self.backupDatabaseDialog.close()

    def deleteTablesData(self):

        logging.info("delete table data")

        # Question when data will be deleted
        reply = QMessageBox.question(self, 'Message',
            "Before erasing all the data. It is recommended that you back up the database.\nAre you sure want to delete all the data?", QMessageBox.Yes, QMessageBox.No)

        if reply == QMessageBox.Yes:
            
            logging.info("Yes! table deleted")

            # delete all data in database table, reset autoincrement, vacuum to clear the data
            qsqlQuery = QSqlQuery()
            qsqlQuery.exec("DELETE FROM Site")
            qsqlQuery.exec("DELETE FROM sqlite_sequence WHERE name = 'Site'")
            qsqlQuery.exec("DELETE FROM HorizonAll")
            qsqlQuery.exec("DELETE FROM sqlite_sequence WHERE name = 'HorizonAll'")
            qsqlQuery.exec("vacuum")

            ## Clear the model
            self.siteDialog.Number_Form_CB.model().clear()
            self.horizonDialog.Number_Form_CB_Hor.model().clear()
            
            # self.horizonModel.clear()
        else:
            # print("No!")
            logging.info("No! delete table canceled")
   
    def aboutSHPy(self):
        
        logging.info("about SH-Py")

        self.aboutDialog = QDialog()
        verticalLayout = QVBoxLayout()
        title = QLabel("Site and Horizon Entry         ")
        version = QLabel("Version: 1.10")
        verticalLayout.addWidget(title)
        verticalLayout.addWidget(version)

        self.aboutDialog.setLayout(verticalLayout)
        self.aboutDialog.setWindowTitle("SHPy ")
        self.aboutDialog.setMaximumWidth(100)
        # self.aboutDialog.setGeometry(500,200,800,100)
        self.aboutDialog.exec()

    def showTable(self):
        """ Show selection of model/database to show  """

        logging.info("Show Table")
        
        # initilize dialog for chossing between site or horizon to show the table of
        self.tableChooseDialog = QDialog()
        # setting the layout
        self.vLayout = QVBoxLayout()
        # add the combobox
        self.tableChooseComboBox = QComboBox()
        # add site and horizon item
        self.tableChooseComboBox.addItem("Site")
        self.tableChooseComboBox.addItem("Horizon")
        # add the pushbutton
        self.tablePushButton = QPushButton("OK")
        # add the combobox and pushButton widgets to the vertical layout
        self.vLayout.addWidget(self.tableChooseComboBox)
        self.vLayout.addWidget(self.tablePushButton)
        # set the layout of the tableChooseDialog to vertical layout (vLayout)
        self.tableChooseDialog.setLayout(self.vLayout)
        # event clicked
        self.tablePushButton.clicked.connect(self.tableChooseClicked)
        # set the geometry dialog window
        self.tableChooseDialog.setGeometry(570, 300, 200, 100)
        # execute the dialog to show the window
        self.tableChooseDialog.exec()
        
    def initRelationalTable(self):
        """ setting the relational database so the code in site/horizon database show
            as English/Indonesia translated """
        # logging.info("init relational database - language")
        logging.warning("init relational database - language")

        self.siteRelationalModel = QSqlRelationalTableModel()
        self.siteRelationalModel.setTable("Site")
        self.siteRelationalModel.setEditStrategy(self.siteRelationalModel.OnManualSubmit)
        self.siteRelationalModel.setRelation(10, QSqlRelation("KindObs", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(16, QSqlRelation("Landform", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(17, QSqlRelation("Relief", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(19, QSqlRelation("SlopeShape", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(20, QSqlRelation("SlopePosition", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(21, QSqlRelation("ParentMaterialAccu", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(22, QSqlRelation("ParentMaterialLitho", "CODE", "TEXT_ENGLISH"))
        # self.siteRelationalModel.setRelation(23, QSqlRelation("ParentMaterialLitho", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(24, QSqlRelation("Drainage", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(25, QSqlRelation("Permeability", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(26, QSqlRelation("Runoff", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(27, QSqlRelation("ErosionType", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(28, QSqlRelation("ErosionDegree", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(30, QSqlRelation("LandCoverMain", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(31, QSqlRelation("LandCoverClimate", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(32, QSqlRelation("LandCoverVegetation", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(33, QSqlRelation("LandCoverPerformance", "CODE", "TEXT_ENGLISH"))
        self.siteRelationalModel.setRelation(34, QSqlRelation("LandCoverMain", "CODE", "TEXT_ENGLISH"))
        # self.siteRelationalModel.setRelation(34, QSqlRelation("LandCoverClimate", "CODE", "TEXT_ENGLISH"))
        # self.siteRelationalModel.setRelation(34, QSqlRelation("LandCoverVegetation", "CODE", "TEXT_ENGLISH"))
        # self.siteRelationalModel.setRelation(34, QSqlRelation("LandCoverPerformance", "CODE", "TEXT_ENGLISH"))


        self.siteRelationalModel.setHeaderData(10, Qt.Horizontal, "Kind of Observation")
        self.siteRelationalModel.setHeaderData(16, Qt.Horizontal, "Landform")
        self.siteRelationalModel.setHeaderData(17, Qt.Horizontal, "Relief")
        self.siteRelationalModel.setHeaderData(19, Qt.Horizontal, "Slope Shape")
        self.siteRelationalModel.setHeaderData(20, Qt.Horizontal, "Slope Position")
        self.siteRelationalModel.setHeaderData(21, Qt.Horizontal, "Accumulation")
        self.siteRelationalModel.setHeaderData(22, Qt.Horizontal, "Lithology 1")
        self.siteRelationalModel.setHeaderData(23, Qt.Horizontal, "Lithology 2")
        self.siteRelationalModel.setHeaderData(24, Qt.Horizontal, "Drainage")
        self.siteRelationalModel.setHeaderData(25, Qt.Horizontal, "Permeability")
        self.siteRelationalModel.setHeaderData(26, Qt.Horizontal, "Runoff")
        self.siteRelationalModel.setHeaderData(27, Qt.Horizontal, "Erosion Type")
        self.siteRelationalModel.setHeaderData(28, Qt.Horizontal, "Erosion Degree")
        self.siteRelationalModel.setHeaderData(30, Qt.Horizontal, "Land Cover Main")
        self.siteRelationalModel.setHeaderData(31, Qt.Horizontal, "Land Cover Climate")
        self.siteRelationalModel.setHeaderData(32, Qt.Horizontal, "Land Cover Vegetation")
        self.siteRelationalModel.setHeaderData(33, Qt.Horizontal, "Land Cover Performance")
        self.siteRelationalModel.setHeaderData(34, Qt.Horizontal, "Land Cover Main 2")
        self.siteRelationalModel.setHeaderData(35, Qt.Horizontal, "Land Cover Climate 2")
        self.siteRelationalModel.setHeaderData(36, Qt.Horizontal, "Land Cover Vegetation 2")
        self.siteRelationalModel.setHeaderData(37, Qt.Horizontal, "Land Cover Performance 2")


        self.siteRelationalModel.select()
    
    def closeEvent(self, event):
        """ setting confirmation of the close of main window """
        logging.info("closed event")
        # print("closed")
        reply = QMessageBox.question(self, 'Message',
            "Are you sure want to quit?", QMessageBox.Yes, QMessageBox.No)

        if reply == QMessageBox.Yes:
            event.accept()
            logging.info("Yes SH-Py windows closed")
        else:
            event.ignore()
            logging.info("No, quit SH-Py windows canceled")

    def reloadClicked(self):
        
        logging.info("reload button clicked")

        # initilize site model
        initializeModel(self.siteModel, "Site")
        # initilize horizon model
        initializeModel(self.horizonModel, "HorizonAll")
        # populate the model required
        self.populateModel()

        ###
        #    SiteWindow.initializeModel(self.siteModel, "Site", '')
        #    self.siteDialog.Number_Form_CB.setModel(self.siteModel)
    
    def horizonClicked(self):
        # logging
        logging.info("Open Horizon Window")
        
        # show the Horizon Window
        # TODO: set the model for Number_Form_CB_Hor, if data is deleted
        self.horizonDialog.show()

        # # initilize site model
        # initializeModel(self.siteModel, "Site")
        # # initilize horizon model
        # initializeModel(self.horizonModel, "HorizonAll")
        # # populate the model required
        # self.populateModel()

    def siteClicked(self):
        # logging
        logging.info("Open Site Window")

        # event to show Site Window 
        self.siteDialog.show()

    def tableChooseClicked(self):
        """ Event dialog showing tables from site or horizon model/database tables """
        logging.info("Table choose clicked")

        # Close the showTable function
        self.tableChooseDialog.close()

        # add list of the two model
        list = [self.siteTableModel, self.horizonTableModel]
        # set the tableView
        self.siteTableView = QTableView()
        self.siteTableView.setModel(list[self.tableChooseComboBox.currentIndex()])
        # set the tableDialog window
        self.tableDialog = QDialog()
        self.vLayout = QVBoxLayout()
        self.vLayout.addWidget(self.siteTableView)
        self.tableDialog.setLayout(self.vLayout)
        self.tableDialog.setWindowTitle("Table")
        self.tableDialog.setGeometry(100, 100, 1200, 600)
        self.tableDialog.exec()
         
    def exportToDocx(self):
        """ Setting the dialog selections and the settings required for export to docx
            Language prefered, range of row/page, and the save file """ 
        # TODO: must add check if the data is empty dont crash the SH-Py!!!   

        logging.info("exporting to Docx format")

        ### Hint or Placeholder for maximum range of number form entry
        list = []
        for i in range(1, self.horizonModel.rowCount()+1):
            proxy = QSortFilterProxyModel()
            proxy.setSourceModel(self.horizonModel)
            # set filter to NoForm column (index 1)
            proxy.setFilterKeyColumn(1)
            # set the filter using the list of the range entered above
            proxy.setFilterFixedString(str(i))
            # matching the value with the source model (site/horizon model)
            matchingIndex = proxy.mapToSource(proxy.index(0,0))
            # get the row of NoForm 
            row = matchingIndex.row()
            list.append(row)
            logging.info("Row: %s - %s",i, row)
        logging.info("list: %s", list)
        maxForm = max(list) + 1
        minForm = min(list) + 1

        ### make dialog and the widget for 'export the docx'
        self.exportDialog = QDialog()
        self.gridLayout = QGridLayout()
        self.languageLabel = QLabel("Language:")
        self.languageComboBox = QComboBox()
        self.languageComboBox.addItem("")
        self.languageComboBox.addItem("INDONESIA")
        self.languageComboBox.addItem("ENGLISH")
        self.languageComboBox.setCurrentIndex(2)
        self.rangeLabel = QLabel("Range Number Form:")
        self.rangeLineEdit = QLineEdit()
        # set the placeholder for the maximum range of number form
        self.rangeLineEdit.setPlaceholderText(str(minForm) + "-" + str(maxForm))
        # set the range automatically to max the number form
        self.rangeLineEdit.setText(str(minForm) + "-" + str(maxForm))
        self.saveDocxFile = QLineEdit()
        self.saveDocxFile.setReadOnly(True)
        self.savePushButton = QPushButton("Docx Location")
        self.okPushButton = QPushButton("OK")
        self.cancelPushButton = QPushButton("Cancel")

        self.gridLayout.setColumnStretch(1, 1)
        self.gridLayout.setColumnMinimumWidth(1, 150)
        ### add the widget to grid layout
        self.gridLayout.addWidget(self.languageLabel, 0, 0)
        self.gridLayout.addWidget(self.languageComboBox, 0, 1)
        self.gridLayout.addWidget(self.rangeLabel, 1, 0)
        self.gridLayout.addWidget(self.rangeLineEdit, 1, 1)
        self.gridLayout.addWidget(self.saveDocxFile, 2, 0, 2, 2)
        self.gridLayout.addWidget(self.savePushButton, 4, 0, 1, 2)
        self.gridLayout.addWidget(self.okPushButton, 5, 0)
        self.gridLayout.addWidget(self.cancelPushButton, 5, 1)
        # set the layout of exportDialog to gridLayout
        self.exportDialog.setLayout(self.gridLayout)
              
        self.okPushButton.clicked.connect(self.okClicked)
        self.cancelPushButton.clicked.connect(self.exportDialog.close)
        self.savePushButton.clicked.connect(self.setSaveDocx)
        self.exportDialog.setWindowTitle("Export to Docx")

        self.exportDialog.exec()
    
    def populateModel(self):
        # numberFormText = self.Number_Form_CB.currentText()
        # row = self.Number_Form_CB.findText(numberFormText)
        logging.info("Populate Model")

        ### Site Model 
        self.kindObsModel = QSqlTableModel()
        self.landFormModel = QSqlTableModel()
        self.reliefModel = QSqlTableModel()
        self.slopeShapeModel = QSqlTableModel()
        self.slopePositionModel = QSqlTableModel()
        self.parentMaterialAccumulationModel = QSqlTableModel()
        self.parentMaterialLithology1Model = QSqlTableModel()
        self.parentMaterialLithology2Model = QSqlTableModel()
        self.drainageClassModel = QSqlTableModel()
        self.permeabilityModel = QSqlTableModel()
        self.runoffModel = QSqlTableModel()
        self.erosionTypeModel = QSqlTableModel()
        self.erosionDegreeModel = QSqlTableModel()
        self.landCoverMain1Model = QSqlTableModel()
        self.landCoverMain2Model = QSqlTableModel()
        self.landCoverClimate1Model = QSqlTableModel()
        self.landCoverClimate2Model = QSqlTableModel()
        self.landCoverVegetation1Model = QSqlTableModel()
        self.landCoverVegetation2Model = QSqlTableModel()
        self.landCoverPerformance1Model = QSqlTableModel()
        self.landCoverPerformance2Model = QSqlTableModel()
        self.soilMoistureRegimeModel = QSqlTableModel()
        self.soilTemperatureRegimeModel = QSqlTableModel()
        self.epipedonModel = QSqlTableModel()
        self.subSurfaceHorizon1Model = QSqlTableModel()
        self.subSurfaceHorizon2Model = QSqlTableModel()
        self.subSurfaceHorizon3Model = QSqlTableModel()
        self.subSurfaceHorizon4Model = QSqlTableModel()
        self.greatGroupModel = QSqlTableModel()
        self.subGroupModel = QSqlTableModel()
        self.pscModel = QSqlTableModel()
        self.mineralModel = QSqlTableModel()
        self.reactionModel = QSqlTableModel()
        self.temperatureModel = QSqlTableModel()

        ### Initialize All SiteWindow Model
        initializeModel(self.kindObsModel, "KindObs")
        initializeModel(self.landFormModel, "Landform")
        initializeModel(self.reliefModel, "Relief")
        initializeModel(self.slopeShapeModel, "SlopeShape")
        initializeModel(self.slopePositionModel, "SlopePosition")
        initializeModel(self.parentMaterialAccumulationModel, "ParentMaterialAccu")
        initializeModel(self.parentMaterialLithology1Model, "ParentMaterialLItho")
        initializeModel(self.parentMaterialLithology2Model, "ParentMaterialLItho")
        initializeModel(self.drainageClassModel, "Drainage")
        initializeModel(self.permeabilityModel, "Permeability")
        initializeModel(self.runoffModel, "Runoff")
        initializeModel(self.erosionTypeModel, "ErosionType")
        initializeModel(self.erosionDegreeModel, "ErosionDegree")
        initializeModel(self.landCoverMain1Model, "LandCoverMain")
        initializeModel(self.landCoverMain2Model, "LandCoverMain")
        initializeModel(self.landCoverClimate1Model, "LandCoverClimate")
        initializeModel(self.landCoverClimate2Model, "LandCoverClimate")
        initializeModel(self.landCoverVegetation1Model, "LandCoverVegetaion")
        initializeModel(self.landCoverVegetation2Model, "LandCoverVegetaion")
        initializeModel(self.landCoverPerformance1Model, "LandCoverPerformance")
        initializeModel(self.landCoverPerformance2Model, "LandCoverPerformance")
        initializeModel(self.soilMoistureRegimeModel, "MoistureRegime")
        initializeModel(self.soilTemperatureRegimeModel, "TemperatureRegime")
        initializeModel(self.epipedonModel, "Epipedon")
        initializeModel(self.subSurfaceHorizon1Model, "SubSurfaceHorizon")
        initializeModel(self.subSurfaceHorizon2Model, "SubSurfaceHorizon")
        initializeModel(self.subSurfaceHorizon3Model, "SubSurfaceHorizon")
        initializeModel(self.subSurfaceHorizon4Model, "SubSurfaceHorizon")
        initializeModel(self.greatGroupModel, "GreatGroup")
        initializeModel(self.subGroupModel, "SubGroup")
        initializeModel(self.pscModel, "PSC")
        initializeModel(self.mineralModel, "Mineral")
        initializeModel(self.reactionModel, "Reaction")
        initializeModel(self.temperatureModel, "TemperatureRegime")


        ### Horizon Model
        self.horBoundaryDistModel = QSqlTableModel()
        self.horBoundaryTopoModel = QSqlTableModel()
        self.munsellModel = QSqlTableModel()
        # self.proxySoilColor1Model = QSortFilterProxyModel()
        # self.proxySoilColor1Model.setSourceModel(self.munsellModel)
        # self.proxySoilColor2Model = QSortFilterProxyModel()
        # self.proxySoilColor2Model.setSourceModel(self.munsellModel)
        # self.proxySoilColor3Model = QSortFilterProxyModel()
        # self.proxySoilColor3Model.setSourceModel(self.munsellModel)
        self.textureClassModel = QSqlTableModel()
        self.textureSandModel = QSqlTableModel()
        self.modSizeModel = QSqlTableModel()
        self.modAbundanceModel = QSqlTableModel()
        self.structureShapeModel = QSqlTableModel()
        self.structureSizeModel = QSqlTableModel()
        self.structureGradeModel = QSqlTableModel()
        self.structureRelationModel = QSqlTableModel()
        # self.proxyStructureShape_2_Model = QSortFilterProxyModel()
        # self.proxyStructureShape_2_Model.setSourceModel(self.structureShape_1_Model)
        # self.proxyStructureSize_2_Model = QSortFilterProxyModel()
        # self.proxyStructureSize_2_Model.setSourceModel(self.structureSize_1_Model)
        # self.proxyStructureGrade_2_Model = QSortFilterProxyModel()
        # self.proxyStructureGrade_2_Model.setSourceModel(self.structureGrade_1_Model)
        self.consistencyMoistModel = QSqlTableModel()
        self.consistencyStickinessModel = QSqlTableModel()
        self.consistencyPlasticityModel = QSqlTableModel()
        self.mottleAbundanceModel = QSqlTableModel()
        self.mottleSizeModel = QSqlTableModel()
        self.mottleContrastModel = QSqlTableModel()
        self.mottleShapeModel = QSqlTableModel()
        self.mottleHueModel = QSqlTableModel()
        self.mottleValueModel = QSqlTableModel()
        self.mottleChromaModel = QSqlTableModel()
        self.rootsFineModel = QSqlTableModel()
        self.rootsMediumModel = QSqlTableModel()
        self.rootsCoarseModel = QSqlTableModel()

        ### Initialize Horizon Model
        initializeModel(self.horBoundaryDistModel, "BoundaryDist")
        initializeModel(self.horBoundaryTopoModel, "BoundaryTopo")
        initializeModel(self.horBoundaryTopoModel, "BoundaryTopo")
        initializeModel(self.munsellModel, "SoilColor")
        initializeModel(self.textureClassModel, "TextureClass")
        initializeModel(self.textureSandModel, "TextureSand")
        initializeModel(self.modSizeModel, "TextureModSize")
        initializeModel(self.modAbundanceModel, "TextureModAbundance")
        initializeModel(self.structureShapeModel, "StructureShape")
        initializeModel(self.structureSizeModel, "StructureSize")
        initializeModel(self.structureGradeModel, "StructureGrade")
        initializeModel(self.structureRelationModel, "StructureRelation")
        initializeModel(self.consistencyMoistModel, "ConsistenceMoist")
        initializeModel(self.consistencyStickinessModel, "ConsistenceStickiness")
        initializeModel(self.consistencyPlasticityModel, "ConsistencePlasticity")
        initializeModel(self.mottleAbundanceModel, "MottlesAbundance")
        initializeModel(self.mottleSizeModel, "MottlesSize")
        initializeModel(self.mottleContrastModel, "MottlesContrast")
        initializeModel(self.mottleShapeModel, "MottlesShape")
        initializeModel(self.rootsFineModel, "RootsFine")
        initializeModel(self.rootsMediumModel, "RootsMedium")
        initializeModel(self.rootsCoarseModel, "RootsCoarse")

    def okClicked(self):
        """ implementation to show the model of site and horizon to microsoft word (docx)
            1st setting the style of the document and then looping through the range of row/page
            desired from the users. then get the data from desired row above if the is in "code/key"
            get the TEXT_INDO(index =1)/TEXT_ENGLISH (index=2) from the related database tables using
            matching function from QSortFilterProxyModel. The data then insert to add_run() method 
            from "python-docx" module. """

        # TODO: 1. add margin and paper settings, (FIXED!) 
        #       2. add placeholder to Number form range from (the lowest to highest/ aka all form) (FIXED!)
        #       3. maybe fixed the export if horizon database is less than site database row
        #          so the output of the all horizon in docx not "NONENONENONENONE"!!! (FIXED!)
        #       4. add documentation in main, site and horizon window. 
        #       5. try add log/info/info in main, site and horizon window using python logger
        #       6. Try to entry some soil survei result!
        #       7. Try fixed any bugs in main, site and horizon window 
        #       8. add new features like export to excel, save database, etc    
        #       9. convert to 'exe' file       
        #       10. Change the Layout/order format of docx file depending on the language selected English or Indonesia
        #       11. Fixed the dynamically height of completer popup tableview (FIXED! using hacks! must find better solution)
        #       13. export database/backup database to another name and folder
        #       14. join/merge database from another database
        #       15. export to excel/csv
        #       16. add messagebox for confirmation in update, insert, delete, etcetera.
        #       17. add copy another number form method for horizon windows (FIXED!)
        #       18. must check if number form in Site Window/Database is less than Horizon Window/Database(Fixed! - I think is not necessary)
        #       19. Bug 1: cant update data in Site Window. Must Check! (FIXED!)
        #       20. Add create new database tables for site and horizon 
        #       21. Add utilities to delete the data from site and horizon database tables
        #       22. After insert button clicked the field is emptied. must check
        #           I think is something with the mapper (FIXED!-using "mapper.toLast()")
        #       23. Make the database query using another thread not main thread
        #       24. Add column createdDate and modifiedDate in Site and Horison Database and implement that
        #           in the operation.
        
        logging.info("Ok exporting to Docx clicked")

        # define the document
        document = Document()
        # Setting the Paper Size (A4)
        section = document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        # setting the page margins
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        # set the styles for access the font and then apply them to paragraph
        style = document.styles['Normal']
        font = style.font
        font.name = "Tahoma"
        font.size = Pt(11)
        
        # check if range edit is not empty
        # TODO: also check if value entered is not numeric/int
        if self.rangeLineEdit.text() != "":
            # looping through the range entered
            for i in parseRangeList(self.rangeLineEdit.text()):
                # set the proxy Site model
                proxy = QSortFilterProxyModel()
                proxy.setSourceModel(self.siteModel)
                # set filter to NoForm column (index 1)
                proxy.setFilterKeyColumn(1)
                # set the filter using the list of the range entered above
                proxy.setFilterFixedString(str(i))
                # matching the value with the source model (site/horizon model)
                matchingIndex = proxy.mapToSource(proxy.index(0,0))
                # get the row of NoForm 
                row = matchingIndex.row()

                # set the proxy Horizon model
                proxyHorizon = QSortFilterProxyModel()
                proxyHorizon.setSourceModel(self.horizonModel)
                # set filter to NoForm column (index 1)
                proxyHorizon.setFilterKeyColumn(1)
                # set the filter using the list of the range entered above
                proxyHorizon.setFilterFixedString(str(i))
                # matching the value with the source model (site/horizon model)
                matchingIndexHorizon = proxyHorizon.mapToSource(proxyHorizon.index(0,0))
                # get the row of NoForm 
                horizonRow = matchingIndexHorizon.row()

                ##
                logging.info("Site row: %s", row)
                logging.info("Horizon row: %s", horizonRow)
                logging.info("Number Form: %s", row)
              


                ### get the data from row above from site model
                numberFormText = str(self.siteModel.data(self.siteModel.index(row,1)))
                sptText = str(self.siteModel.data(self.siteModel.index(row,2)))
                provinsiText = str(self.siteModel.data(self.siteModel.index(row,3)))
                kabupatenText = str(self.siteModel.data(self.siteModel.index(row,4)))
                kecamatanText = str(self.siteModel.data(self.siteModel.index(row,5)))
                desaText = str(self.siteModel.data(self.siteModel.index(row,6)))
                dateText = str(self.siteModel.data(self.siteModel.index(row,7)))
                initialNameText = str(self.siteModel.data(self.siteModel.index(row,8)))
                observationNumberText = str(self.siteModel.data(self.siteModel.index(row,9)))
                kindObservationText = str(self.siteModel.data(self.siteModel.index(row,10)))
                utmZone1Text =str(self.siteModel.data(self.siteModel.index(row,11)))
                utmZone2Text = str(self.siteModel.data(self.siteModel.index(row,12)))
                xEastText = str(self.siteModel.data(self.siteModel.index(row,13)))
                yNorthText = str(self.siteModel.data(self.siteModel.index(row,14)))
                elevationText = str(self.siteModel.data(self.siteModel.index(row,15)))
                landformText = str(self.siteModel.data(self.siteModel.index(row,16)))
                reliefText = str(self.siteModel.data(self.siteModel.index(row,17)))
                slopePercentageText = str(self.siteModel.data(self.siteModel.index(row,18)))
                slopeShapeText = str(self.siteModel.data(self.siteModel.index(row,19)))
                slopePositionText = str(self.siteModel.data(self.siteModel.index(row,20)))
                parentMaterialAccumulationText = str(self.siteModel.data(self.siteModel.index(row,21)))
                lithology1Text = str(self.siteModel.data(self.siteModel.index(row,22)))
                lithology2Text = str(self.siteModel.data(self.siteModel.index(row,23)))
                # print("lithology2Text:", lithology2Text)
                drainageClassText = str(self.siteModel.data(self.siteModel.index(row,24)))
                permeabilityClassText = str(self.siteModel.data(self.siteModel.index(row,25)))
                runoffText = str(self.siteModel.data(self.siteModel.index(row,26)))
                erosionTypeText = str(self.siteModel.data(self.siteModel.index(row,27)))
                erosionDegreeText = str(self.siteModel.data(self.siteModel.index(row,28)))
                effectiveSoilDepthText = str(self.siteModel.data(self.siteModel.index(row,29)))
                landCoverMain1Text = str(self.siteModel.data(self.siteModel.index(row,30)))
                landCoverClimate1Text = str(self.siteModel.data(self.siteModel.index(row,31)))
                landCoverVegetation1Text = str(self.siteModel.data(self.siteModel.index(row,32)))
                landCoverPerformance1Text = str(self.siteModel.data(self.siteModel.index(row,33)))
                landCoverMain2Text = str(self.siteModel.data(self.siteModel.index(row,34)))
                landCoverClimate2Text = str(self.siteModel.data(self.siteModel.index(row,35)))
                landCoverVegetation2Text = str(self.siteModel.data(self.siteModel.index(row,36)))
                landCoverPerformance2Text = str(self.siteModel.data(self.siteModel.index(row,37)))
                epipedonText = str(self.siteModel.data(self.siteModel.index(row,38)))
                upperLimit1Text = str(self.siteModel.data(self.siteModel.index(row,39)))
                lowerLimit1Text = str(self.siteModel.data(self.siteModel.index(row,40)))
                subSurfaceHorizon1Text = str(self.siteModel.data(self.siteModel.index(row,41)))
                upperLimit2Text = str(self.siteModel.data(self.siteModel.index(row,42)))
                lowerLimit2Text = str(self.siteModel.data(self.siteModel.index(row,43)))
                subSurfaceHorizon2Text = str(self.siteModel.data(self.siteModel.index(row,44)))
                upperLimit3Text = str(self.siteModel.data(self.siteModel.index(row,45)))
                lowerLimit3Text = str(self.siteModel.data(self.siteModel.index(row,46)))
                subSurfaceHorizon3Text = str(self.siteModel.data(self.siteModel.index(row,47)))
                upperLimit4Text = str(self.siteModel.data(self.siteModel.index(row,48)))
                lowerLimit4Text = str(self.siteModel.data(self.siteModel.index(row,49)))
                subSurfaceHorizon4Text = str(self.siteModel.data(self.siteModel.index(row,50)))
                upperLimit5Text = str(self.siteModel.data(self.siteModel.index(row,51)))
                lowerLimit5Text = str(self.siteModel.data(self.siteModel.index(row,52)))
                soilMoistureRegimeText = str(self.siteModel.data(self.siteModel.index(row,53)))
                soilTemperatureRegimeText = str(self.siteModel.data(self.siteModel.index(row,54)))
                taxonomyYearText = str(self.siteModel.data(self.siteModel.index(row,55)))
                greatGroupText = str(self.siteModel.data(self.siteModel.index(row,56)))
                subGroupText = str(self.siteModel.data(self.siteModel.index(row,57)))
                pscText = str(self.siteModel.data(self.siteModel.index(row,58)))
                mineralText = str(self.siteModel.data(self.siteModel.index(row,59)))
                reactionText = str(self.siteModel.data(self.siteModel.index(row,60)))
                temperatureText = str(self.siteModel.data(self.siteModel.index(row,61)))
                
                ### Horizon 1
                HorDesignDisconText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 2))))
                HorDesignMasterText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 3))))
                HorDesignSubText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 4))))
                HorDesignNumberText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 5))))
                HorUpperFromText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 6))))
                HorUpperToText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 7))))
                HorLowerFromText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 8))))
                HorLowerToText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 9))))
                BoundaryDistText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 10))))
                BoundaryTopoText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 11))))
                SoilColorHue_1_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 13))))
                SoilColorValue_1_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 14))))
                SoilColorChroma_1_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 15))))
                SoilColorHue_2_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 17))))
                SoilColorValue_2_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 18))))
                SoilColorChroma_2_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 19))))
                SoilColorHue_3_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 21))))
                SoilColorValue_3_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 22))))
                SoilColorChroma_3_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 23))))
                TextureClassText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 24))))
                TextureSandText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 25))))
                TextureModSizeText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 26))))
                TextureModAbundanceText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 27))))
                StructureShape_1_Text_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 28))))
                StructureSize_1_Text_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 29))))
                StructureGrade_1_Text_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 30))))
                StructureRelationText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 31))))
                StructureShape_2_Text_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 32))))
                StructureSize_2_Text_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 33))))
                StructureGrade_2_Text_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 34))))
                ConsistenceMoistText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 35))))
                ConsistenceStickinessText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 36))))
                ConsistencePlasticityText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 37))))
                MottleAbundance_1_Text_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 38))))
                MottleSize_1_Text_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 39))))
                MottleContrast_1_Text_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 40))))
                MottleShape_1_Text_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 41))))
                MottleHue_1_Text_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 42))))
                MottleValue_1_Text_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 43))))
                MottleChroma_1_Text_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 44))))
                RootFineText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 45))))
                RootMediumText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 46))))
                RootCoarseText_Hor_1 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 47)))) 
                
                ### Horizon 2
                HorDesignDisconText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 48))))
                HorDesignMasterText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 49))))
                HorDesignSubText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 50))))
                HorDesignNumberText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 51))))
                HorUpperFromText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 52))))
                HorUpperToText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 53))))
                HorLowerFromText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 54))))
                HorLowerToText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 55))))
                BoundaryDistText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 56))))
                BoundaryTopoText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 57))))
                SoilColorHue_1_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 59))))
                SoilColorValue_1_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 60))))
                SoilColorChroma_1_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 61))))
                SoilColorHue_2_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 63))))
                SoilColorValue_2_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 64))))
                SoilColorChroma_2_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 65))))
                SoilColorHue_3_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 67))))
                SoilColorValue_3_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 68))))
                SoilColorChroma_3_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 69))))
                TextureClassText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 70))))
                TextureSandText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 71))))
                TextureModSizeText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 72))))
                TextureModAbundanceText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 73))))
                StructureShape_1_Text_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 74))))
                StructureSize_1_Text_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 75))))
                StructureGrade_1_Text_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 76))))
                StructureRelationText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 77))))
                StructureShape_2_Text_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 78))))
                StructureSize_2_Text_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 79))))
                StructureGrade_2_Text_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 80))))
                ConsistenceMoistText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 81))))
                ConsistenceStickinessText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 82))))
                ConsistencePlasticityText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 83))))
                MottleAbundance_1_Text_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 84))))
                MottleSize_1_Text_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 85))))
                MottleContrast_1_Text_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 86))))
                MottleShape_1_Text_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 87))))
                MottleHue_1_Text_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 88))))
                MottleValue_1_Text_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 89))))
                MottleChroma_1_Text_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 90))))
                RootFineText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 91))))
                RootMediumText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 92))))
                RootCoarseText_Hor_2 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 93)))) 
                
                ### Horizon 3
                HorDesignDisconText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 94))))
                HorDesignMasterText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 95))))
                HorDesignSubText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 96))))
                HorDesignNumberText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 97))))
                HorUpperFromText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 98))))
                HorUpperToText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 99))))
                HorLowerFromText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 100))))
                HorLowerToText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 101))))
                BoundaryDistText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 102))))
                BoundaryTopoText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 103))))
                SoilColorHue_1_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 105))))
                SoilColorValue_1_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 106))))
                SoilColorChroma_1_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 107))))
                SoilColorHue_2_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 109))))
                SoilColorValue_2_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 110))))
                SoilColorChroma_2_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 111))))
                SoilColorHue_3_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 113))))
                SoilColorValue_3_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 114))))
                SoilColorChroma_3_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 115))))
                TextureClassText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 116))))
                TextureSandText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 117))))
                TextureModSizeText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 118))))
                TextureModAbundanceText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 119))))
                StructureShape_1_Text_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 120))))
                StructureSize_1_Text_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 121))))
                StructureGrade_1_Text_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 122))))
                StructureRelationText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 123))))
                StructureShape_2_Text_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 124))))
                StructureSize_2_Text_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 125))))
                StructureGrade_2_Text_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 126))))
                ConsistenceMoistText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 127))))
                ConsistenceStickinessText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 128))))
                ConsistencePlasticityText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 129))))
                MottleAbundance_1_Text_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 130))))
                MottleSize_1_Text_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 131))))
                MottleContrast_1_Text_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 132))))
                MottleShape_1_Text_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 133))))
                MottleHue_1_Text_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 134))))
                MottleValue_1_Text_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 135))))
                MottleChroma_1_Text_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 136))))
                RootFineText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 137))))
                RootMediumText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 138))))
                RootCoarseText_Hor_3 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 139)))) 
                
                ### Horizon 4
                HorDesignDisconText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 140))))
                HorDesignMasterText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 141))))
                HorDesignSubText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 142))))
                HorDesignNumberText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 143))))
                HorUpperFromText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 144))))
                HorUpperToText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 145))))
                HorLowerFromText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 146))))
                HorLowerToText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 147))))
                BoundaryDistText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 148))))
                BoundaryTopoText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 149))))
                SoilColorHue_1_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 151))))
                SoilColorValue_1_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 152))))
                SoilColorChroma_1_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 153))))
                SoilColorHue_2_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 155))))
                SoilColorValue_2_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 156))))
                SoilColorChroma_2_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 157))))
                SoilColorHue_3_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 159))))
                SoilColorValue_3_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 160))))
                SoilColorChroma_3_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 161))))
                TextureClassText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 162))))
                TextureSandText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 163))))
                TextureModSizeText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 164))))
                TextureModAbundanceText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 165))))
                StructureShape_1_Text_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 166))))
                StructureSize_1_Text_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 167))))
                StructureGrade_1_Text_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 168))))
                StructureRelationText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 169))))
                StructureShape_2_Text_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 170))))
                StructureSize_2_Text_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 171))))
                StructureGrade_2_Text_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 172))))
                ConsistenceMoistText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 173))))
                ConsistenceStickinessText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 174))))
                ConsistencePlasticityText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 175))))
                MottleAbundance_1_Text_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 176))))
                MottleSize_1_Text_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 177))))
                MottleContrast_1_Text_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 178))))
                MottleShape_1_Text_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 179))))
                MottleHue_1_Text_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 180))))
                MottleValue_1_Text_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 181))))
                MottleChroma_1_Text_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 182))))
                RootFineText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 183))))
                RootMediumText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 184))))
                RootCoarseText_Hor_4 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 185)))) 
                
                ### Horizon 5
                HorDesignDisconText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 186))))
                HorDesignMasterText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 187))))
                HorDesignSubText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 188))))
                HorDesignNumberText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 189))))
                HorUpperFromText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 190))))
                HorUpperToText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 191))))
                HorLowerFromText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 192))))
                HorLowerToText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 193))))
                BoundaryDistText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 194))))
                BoundaryTopoText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 195))))
                SoilColorHue_1_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 197))))
                SoilColorValue_1_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 198))))
                SoilColorChroma_1_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 199))))
                SoilColorHue_2_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 201))))
                SoilColorValue_2_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 202))))
                SoilColorChroma_2_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 203))))
                SoilColorHue_3_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 205))))
                SoilColorValue_3_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 206))))
                SoilColorChroma_3_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 207))))
                TextureClassText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 208))))
                TextureSandText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 209))))
                TextureModSizeText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 210))))
                TextureModAbundanceText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 211))))
                StructureShape_1_Text_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 212))))
                StructureSize_1_Text_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 213))))
                StructureGrade_1_Text_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 214))))
                StructureRelationText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 215))))
                StructureShape_2_Text_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 216))))
                StructureSize_2_Text_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 217))))
                StructureGrade_2_Text_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 218))))
                ConsistenceMoistText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 219))))
                ConsistenceStickinessText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 220))))
                ConsistencePlasticityText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 221))))
                MottleAbundance_1_Text_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 222))))
                MottleSize_1_Text_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 223))))
                MottleContrast_1_Text_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 224))))
                MottleShape_1_Text_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 225))))
                MottleHue_1_Text_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 226))))
                MottleValue_1_Text_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 227))))
                MottleChroma_1_Text_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 228))))
                RootFineText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 229))))
                RootMediumText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 230))))
                RootCoarseText_Hor_5 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 231)))) 
                
                ### Horizon 6
                HorDesignDisconText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 232))))
                HorDesignMasterText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 233))))
                HorDesignSubText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 234))))
                HorDesignNumberText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 235))))
                HorUpperFromText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 236))))
                HorUpperToText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 237))))
                HorLowerFromText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 238))))
                HorLowerToText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 239))))
                BoundaryDistText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 240))))
                BoundaryTopoText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 241))))
                SoilColorHue_1_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 243))))
                SoilColorValue_1_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 244))))
                SoilColorChroma_1_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 245))))
                SoilColorHue_2_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 247))))
                SoilColorValue_2_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 248))))
                SoilColorChroma_2_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 249))))
                SoilColorHue_3_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 251))))
                SoilColorValue_3_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 252))))
                SoilColorChroma_3_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 253))))
                TextureClassText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 254))))
                TextureSandText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 255))))
                TextureModSizeText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 256))))
                TextureModAbundanceText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 257))))
                StructureShape_1_Text_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 258))))
                StructureSize_1_Text_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 259))))
                StructureGrade_1_Text_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 260))))
                StructureRelationText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 261))))
                StructureShape_2_Text_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 262))))
                StructureSize_2_Text_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 263))))
                StructureGrade_2_Text_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 264))))
                ConsistenceMoistText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 265))))
                ConsistenceStickinessText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 266))))
                ConsistencePlasticityText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 267))))
                MottleAbundance_1_Text_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 268))))
                MottleSize_1_Text_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 269))))
                MottleContrast_1_Text_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 270))))
                MottleShape_1_Text_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 271))))
                MottleHue_1_Text_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 272))))
                MottleValue_1_Text_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 273))))
                MottleChroma_1_Text_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 274))))
                RootFineText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 275))))
                RootMediumText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 276))))
                RootCoarseText_Hor_6 = (str(self.horizonModel.data(self.horizonModel.index(horizonRow, 277)))) 

                ### get the language desirable for this using TEXT_ENGLISH
                ### some of these i thing not all required, can just get the above text
                numberFormTextEnglish = numberFormText 
                sptTextEnglish = sptText 
                provinsiTextEnglish = provinsiText 
                kabupatenTextEnglish = kabupatenText 
                kecamatanTextEnglish = kecamatanText 
                desaTextEnglish = desaText 
                dateTextEnglish = dateText 
                initialNameTextEnglish = initialNameText 
                observationNumberTextEnglish = observationNumberText 
                kindObservationTextEnglish = proxyModelRowIndex(self.kindObsModel, kindObservationText, self.languageComboBox.currentIndex())
                utmZone1TextEnglish = utmZone1Text 
                utmZone2TextEnglish = utmZone2Text 
                xEastTextEnglish = xEastText 
                yNorthTextEnglish = yNorthText 
                elevationTextEnglish = elevationText 
                landformTextEnglish = proxyModelRowIndex(self.landFormModel, landformText, self.languageComboBox.currentIndex())
                reliefTextEnglish = proxyModelRowIndex(self.reliefModel, reliefText, self.languageComboBox.currentIndex())
                slopePercentageTextEnglish = slopePercentageText
                slopeShapeTextEnglish = proxyModelRowIndex(self.slopeShapeModel, slopeShapeText, self.languageComboBox.currentIndex())
                slopePositionTextEnglish = proxyModelRowIndex(self.slopePositionModel, slopePositionText, self.languageComboBox.currentIndex())
                parentMaterialAccumulationTextEnglish = proxyModelRowIndex(self.parentMaterialAccumulationModel, parentMaterialAccumulationText, self.languageComboBox.currentIndex())
                lithology1TextEnglish = proxyModelRowIndex(self.parentMaterialLithology1Model, lithology1Text, self.languageComboBox.currentIndex())
                lithology2TextEnglish = proxyModelRowIndex(self.parentMaterialLithology2Model, lithology2Text, self.languageComboBox.currentIndex()) 
                # print("lithology2TextEnglish", lithology2TextEnglish )
                drainageClassTextEnglish = proxyModelRowIndex(self.drainageClassModel, drainageClassText)
                permeabilityClassTextEnglish = proxyModelRowIndex(self.permeabilityModel, permeabilityClassText, self.languageComboBox.currentIndex())
                runoffTextEnglish = proxyModelRowIndex(self.runoffModel, runoffText, self.languageComboBox.currentIndex())
                erosionTypeTextEnglish = proxyModelRowIndex(self.erosionTypeModel, erosionTypeText, self.languageComboBox.currentIndex())
                erosionDegreeTextEnglish = proxyModelRowIndex(self.erosionDegreeModel, erosionDegreeText, self.languageComboBox.currentIndex())
                effectiveSoilDepthTextEnglish = effectiveSoilDepthText
                landCoverMain1TextEnglish = proxyModelRowIndex(self.landCoverMain1Model, landCoverMain1Text, self.languageComboBox.currentIndex())
                lancCoverClimate1TextEnglish = proxyModelRowIndex(self.landCoverClimate1Model, landCoverClimate1Text, self.languageComboBox.currentIndex())
                landCoverVegetation1TextEnglish = proxyModelRowIndex(self.landCoverVegetation1Model, landCoverVegetation1Text, self.languageComboBox.currentIndex())
                landCoverPerformance1TextEnglish = proxyModelRowIndex(self.landCoverPerformance1Model, landCoverPerformance1Text, self.languageComboBox.currentIndex())
                landCoverMain2TextEnglish = proxyModelRowIndex(self.landCoverMain2Model, landCoverMain2Text, self.languageComboBox.currentIndex())
                lancCoverClimate2TextEnglish = proxyModelRowIndex(self.landCoverClimate2Model, landCoverClimate2Text, self.languageComboBox.currentIndex())
                landCoverVegetation2TextEnglish = proxyModelRowIndex(self.landCoverVegetation2Model, landCoverVegetation2Text, self.languageComboBox.currentIndex())
                landCoverPerformance2TextEnglish = proxyModelRowIndex(self.landCoverPerformance2Model, landCoverPerformance2Text, self.languageComboBox.currentIndex())
                epipedonTextEnglish = proxyModelRowIndex(self.epipedonModel, epipedonText, self.languageComboBox.currentIndex())
                upperLimit1TextEnglish = upperLimit1Text
                lowerLimit1TextEnglish = lowerLimit1Text
                subSurfaceHorizon1TextEnglish = proxyModelRowIndex(self.subSurfaceHorizon1Model, subSurfaceHorizon1Text, self.languageComboBox.currentIndex())
                upperLimit2TextEnglish = upperLimit2Text
                lowerLimit2TextEnglish = lowerLimit2Text
                subSurfaceHorizon2TextEnglish = proxyModelRowIndex(self.subSurfaceHorizon2Model, subSurfaceHorizon2Text, self.languageComboBox.currentIndex())
                upperLimit3TextEnglish = upperLimit3Text 
                lowerLimit3TextEnglish = lowerLimit3Text
                subSurfaceHorizon3TextEnglish = proxyModelRowIndex(self.subSurfaceHorizon3Model, subSurfaceHorizon3Text, self.languageComboBox.currentIndex())
                upperLimit4TextEnglish = upperLimit4Text
                lowerLimit4TextEnglish = lowerLimit4Text
                subSurfaceHorizon4TextEnglish = proxyModelRowIndex(self.subSurfaceHorizon4Model, subSurfaceHorizon4Text, self.languageComboBox.currentIndex())
                upperLimit5TextEnglish = upperLimit5Text
                lowerLimit5TextEnglish = lowerLimit5Text
                soilMoistureRegimeTextEnglish = proxyModelRowIndex(self.soilMoistureRegimeModel, soilMoistureRegimeText, self.languageComboBox.currentIndex())
                soilTemperatureRegimeTextEnglish = proxyModelRowIndex(self.soilTemperatureRegimeModel, soilTemperatureRegimeText, self.languageComboBox.currentIndex())
                taxonomyYearTextEnglish = taxonomyYearText
                greatGroupTextEnglish = proxyModelRowIndex(self.greatGroupModel, greatGroupText, self.languageComboBox.currentIndex())
                subGroupTextEnglish = proxyModelRowIndex(self.subGroupModel, subGroupText, self.languageComboBox.currentIndex())
                pscTextEnglish = proxyModelRowIndex(self.pscModel, pscText, self.languageComboBox.currentIndex())
                mineralTextEnglish = proxyModelRowIndex(self.mineralModel, mineralText, self.languageComboBox.currentIndex())
                reactionTextEnglish = proxyModelRowIndex(self.reactionModel, reactionText, self.languageComboBox.currentIndex())
                temperatureTextEnglish = proxyModelRowIndex(self.temperatureModel, temperatureText, self.languageComboBox.currentIndex())
                
                # logging.info("mineral text english: %s", lithology2TextEnglish)
                
                ### HORIZON
                ## Horizon 1
                HorDesignDisconText_English_Hor_1 = HorDesignDisconText_Hor_1 
                HorDesignMasterText_English_Hor_1 = HorDesignMasterText_Hor_1 
                HorDesignSubText_English_Hor_1 = HorDesignSubText_Hor_1 
                HorDesignNumberText_English_Hor_1 = HorDesignNumberText_Hor_1 
                HorUpperFromText_English_Hor_1 = HorUpperFromText_Hor_1 
                HorUpperToText_English_Hor_1 = HorUpperToText_Hor_1 
                HorLowerFromText_English_Hor_1 = HorLowerFromText_Hor_1 
                HorLowerToText_English_Hor_1 = HorLowerToText_Hor_1 
                BoundaryDistText_English_Hor_1 = proxyModelRowIndex(self.horBoundaryDistModel, BoundaryDistText_Hor_1, self.languageComboBox.currentIndex())
                BoundaryTopoText_English_Hor_1 = proxyModelRowIndex(self.horBoundaryTopoModel, BoundaryTopoText_Hor_1, self.languageComboBox.currentIndex())
                SoilColor_1_Code_Hor_1 = (SoilColorHue_1_Hor_1 + SoilColorValue_1_Hor_1 + "/"+SoilColorChroma_1_Hor_1)
                SoilColor_1_Text_English_Hor_1 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_1_Code_Hor_1, self.languageComboBox.currentIndex() + 4)
                SoilColor_2_Code_Hor_1 = (SoilColorHue_2_Hor_1 + SoilColorValue_2_Hor_1 + "/"+SoilColorChroma_2_Hor_1)
                SoilColor_2_Text_English_Hor_1 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_2_Code_Hor_1, self.languageComboBox.currentIndex() + 4)
                SoilColor_3_Code_Hor_1 = (SoilColorHue_3_Hor_1 + SoilColorValue_3_Hor_1 + "/"+SoilColorChroma_3_Hor_1)
                SoilColor_3_Text_English_Hor_1 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_3_Code_Hor_1, self.languageComboBox.currentIndex() + 4)  
                TextureClassText_English_Hor_1 = proxyModelRowIndexRegX(self.textureClassModel, TextureClassText_Hor_1, self.languageComboBox.currentIndex())
                TextureSandText_English_Hor_1 = proxyModelRowIndexRegX(self.textureSandModel, TextureSandText_Hor_1, self.languageComboBox.currentIndex())
                TextureModSizeText_English_Hor_1 = proxyModelRowIndex(self.modSizeModel, TextureModSizeText_Hor_1, self.languageComboBox.currentIndex())
                TextureModAbundanceText_English_Hor_1 = proxyModelRowIndex(self.modAbundanceModel, TextureModAbundanceText_Hor_1, self.languageComboBox.currentIndex())
                StructureShape_1_Text_English_Hor_1 = proxyModelRowIndexRegX(self.structureShapeModel, StructureShape_1_Text_Hor_1, self.languageComboBox.currentIndex())
                StructureSize_1_Text_English_Hor_1 = proxyModelRowIndexRegX(self.structureSizeModel, StructureSize_1_Text_Hor_1, self.languageComboBox.currentIndex())
                StructureGrade_1_Text_English_Hor_1 =  proxyModelRowIndex(self.structureGradeModel, StructureGrade_1_Text_Hor_1, self.languageComboBox.currentIndex())
                StructureRelationText_English_Hor_1 = proxyModelRowIndex(self.structureRelationModel, StructureRelationText_Hor_1, self.languageComboBox.currentIndex())
                StructureShape_2_Text_English_Hor_1 = proxyModelRowIndexRegX(self.structureShapeModel, StructureShape_2_Text_Hor_1, self.languageComboBox.currentIndex()) 
                StructureSize_2_Text_English_Hor_1 = proxyModelRowIndexRegX(self.structureSizeModel, StructureSize_2_Text_Hor_1, self.languageComboBox.currentIndex())
                StructureGrade_2_Text_English_Hor_1 = proxyModelRowIndex(self.structureGradeModel, StructureGrade_2_Text_Hor_1, self.languageComboBox.currentIndex())
                ConsistenceMoistText_English_Hor_1 = proxyModelRowIndexRegX(self.consistencyMoistModel, ConsistenceMoistText_Hor_1, self.languageComboBox.currentIndex())
                ConsistenceStickinessText_English_Hor_1 = proxyModelRowIndexRegX(self.consistencyStickinessModel, ConsistenceStickinessText_Hor_1, self.languageComboBox.currentIndex())
                ConsistencePlasticityText_English_Hor_1 = proxyModelRowIndexRegX(self.consistencyPlasticityModel, ConsistencePlasticityText_Hor_1, self.languageComboBox.currentIndex())
                MottleAbundance_1_Text_English_Hor_1 = proxyModelRowIndex(self.mottleAbundanceModel, MottleAbundance_1_Text_Hor_1, self.languageComboBox.currentIndex())
                MottleSize_1_Text_English_Hor_1 = proxyModelRowIndex(self.mottleSizeModel, MottleSize_1_Text_Hor_1, self.languageComboBox.currentIndex())
                MottleContrast_1_Text_English_Hor_1 = proxyModelRowIndex(self.mottleContrastModel, MottleContrast_1_Text_Hor_1, self.languageComboBox.currentIndex())
                MottleShape_1_Text_English_Hor_1 = proxyModelRowIndex(self.mottleShapeModel, MottleShape_1_Text_Hor_1, self.languageComboBox.currentIndex())
                MottleColor_1_Code_Hor_1 = (MottleHue_1_Text_Hor_1 + MottleValue_1_Text_Hor_1 + "/" +MottleChroma_1_Text_Hor_1)
                MottleColor_1_Text_English_Hor_1 = proxyModelRowIndexRegXMunsell(self.munsellModel, MottleColor_1_Code_Hor_1, self.languageComboBox.currentIndex() + 4)
                RootFineText_English_Hor_1 = proxyModelRowIndex(self.rootsFineModel, RootFineText_Hor_1, self.languageComboBox.currentIndex())
                RootMediumText_English_Hor_1 = proxyModelRowIndex(self.rootsMediumModel, RootMediumText_Hor_1, self.languageComboBox.currentIndex())
                RootCoarseText_English_Hor_1 = proxyModelRowIndex(self.rootsCoarseModel, RootCoarseText_Hor_1, self.languageComboBox.currentIndex())

                ## Horizon 2
                HorDesignDisconText_English_Hor_2 = HorDesignDisconText_Hor_2 
                HorDesignMasterText_English_Hor_2 = HorDesignMasterText_Hor_2 
                HorDesignSubText_English_Hor_2 = HorDesignSubText_Hor_2 
                HorDesignNumberText_English_Hor_2 = HorDesignNumberText_Hor_2 
                HorUpperFromText_English_Hor_2 = HorUpperFromText_Hor_2 
                HorUpperToText_English_Hor_2 = HorUpperToText_Hor_2 
                HorLowerFromText_English_Hor_2 = HorLowerFromText_Hor_2 
                HorLowerToText_English_Hor_2 = HorLowerToText_Hor_2 
                BoundaryDistText_English_Hor_2 = proxyModelRowIndex(self.horBoundaryDistModel, BoundaryDistText_Hor_2, self.languageComboBox.currentIndex())
                BoundaryTopoText_English_Hor_2 = proxyModelRowIndex(self.horBoundaryTopoModel, BoundaryTopoText_Hor_2, self.languageComboBox.currentIndex())
                SoilColor_1_Code_Hor_2 = (SoilColorHue_1_Hor_2 + SoilColorValue_1_Hor_2 + "/"+SoilColorChroma_1_Hor_2)
                SoilColor_1_Text_English_Hor_2 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_1_Code_Hor_2, self.languageComboBox.currentIndex() + 4)
                SoilColor_2_Code_Hor_2 = (SoilColorHue_2_Hor_2 + SoilColorValue_2_Hor_2 + "/"+SoilColorChroma_2_Hor_2)
                SoilColor_2_Text_English_Hor_2 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_2_Code_Hor_2, self.languageComboBox.currentIndex() + 4)
                SoilColor_3_Code_Hor_2 = (SoilColorHue_3_Hor_2 + SoilColorValue_3_Hor_2 + "/"+SoilColorChroma_3_Hor_2)
                SoilColor_3_Text_English_Hor_2 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_3_Code_Hor_2, self.languageComboBox.currentIndex() + 4)  
                TextureClassText_English_Hor_2 = proxyModelRowIndexRegX(self.textureClassModel, TextureClassText_Hor_2, self.languageComboBox.currentIndex())
                TextureSandText_English_Hor_2 = proxyModelRowIndexRegX(self.textureSandModel, TextureSandText_Hor_2, self.languageComboBox.currentIndex())
                TextureModSizeText_English_Hor_2 = proxyModelRowIndex(self.modSizeModel, TextureModSizeText_Hor_2, self.languageComboBox.currentIndex())
                TextureModAbundanceText_English_Hor_2 = proxyModelRowIndex(self.modAbundanceModel, TextureModAbundanceText_Hor_2, self.languageComboBox.currentIndex())
                StructureShape_1_Text_English_Hor_2 = proxyModelRowIndexRegX(self.structureShapeModel, StructureShape_1_Text_Hor_2, self.languageComboBox.currentIndex())
                StructureSize_1_Text_English_Hor_2 = proxyModelRowIndexRegX(self.structureSizeModel, StructureSize_1_Text_Hor_2, self.languageComboBox.currentIndex())
                StructureGrade_1_Text_English_Hor_2 =  proxyModelRowIndex(self.structureGradeModel, StructureGrade_1_Text_Hor_2, self.languageComboBox.currentIndex())
                StructureRelationText_English_Hor_2 = proxyModelRowIndex(self.structureRelationModel, StructureRelationText_Hor_2, self.languageComboBox.currentIndex())
                StructureShape_2_Text_English_Hor_2 = proxyModelRowIndexRegX(self.structureShapeModel, StructureShape_2_Text_Hor_2, self.languageComboBox.currentIndex()) 
                StructureSize_2_Text_English_Hor_2 = proxyModelRowIndexRegX(self.structureSizeModel, StructureSize_2_Text_Hor_2, self.languageComboBox.currentIndex())
                StructureGrade_2_Text_English_Hor_2 = proxyModelRowIndex(self.structureGradeModel, StructureGrade_2_Text_Hor_2, self.languageComboBox.currentIndex())
                ConsistenceMoistText_English_Hor_2 = proxyModelRowIndexRegX(self.consistencyMoistModel, ConsistenceMoistText_Hor_2, self.languageComboBox.currentIndex())
                ConsistenceStickinessText_English_Hor_2 = proxyModelRowIndexRegX(self.consistencyStickinessModel, ConsistenceStickinessText_Hor_2, self.languageComboBox.currentIndex())
                ConsistencePlasticityText_English_Hor_2 = proxyModelRowIndexRegX(self.consistencyPlasticityModel, ConsistencePlasticityText_Hor_2, self.languageComboBox.currentIndex())
                MottleAbundance_1_Text_English_Hor_2 = proxyModelRowIndex(self.mottleAbundanceModel, MottleAbundance_1_Text_Hor_2, self.languageComboBox.currentIndex())
                MottleSize_1_Text_English_Hor_2 = proxyModelRowIndex(self.mottleSizeModel, MottleSize_1_Text_Hor_2, self.languageComboBox.currentIndex())
                MottleContrast_1_Text_English_Hor_2 = proxyModelRowIndex(self.mottleContrastModel, MottleContrast_1_Text_Hor_2, self.languageComboBox.currentIndex())
                MottleShape_1_Text_English_Hor_2 = proxyModelRowIndex(self.mottleShapeModel, MottleShape_1_Text_Hor_2, self.languageComboBox.currentIndex())
                MottleColor_1_Code_Hor_2 = (MottleHue_1_Text_Hor_2 + MottleValue_1_Text_Hor_2 + "/" +MottleChroma_1_Text_Hor_2)
                MottleColor_1_Text_English_Hor_2 = proxyModelRowIndexRegXMunsell(self.munsellModel, MottleColor_1_Code_Hor_2, self.languageComboBox.currentIndex() + 4)
                RootFineText_English_Hor_2 = proxyModelRowIndex(self.rootsFineModel, RootFineText_Hor_2, self.languageComboBox.currentIndex())
                RootMediumText_English_Hor_2 = proxyModelRowIndex(self.rootsMediumModel, RootMediumText_Hor_2, self.languageComboBox.currentIndex())
                RootCoarseText_English_Hor_2 = proxyModelRowIndex(self.rootsCoarseModel, RootCoarseText_Hor_2, self.languageComboBox.currentIndex())

                ## Horizon 3
                HorDesignDisconText_English_Hor_3 = HorDesignDisconText_Hor_3 
                HorDesignMasterText_English_Hor_3 = HorDesignMasterText_Hor_3 
                HorDesignSubText_English_Hor_3 = HorDesignSubText_Hor_3 
                HorDesignNumberText_English_Hor_3 = HorDesignNumberText_Hor_3 
                HorUpperFromText_English_Hor_3 = HorUpperFromText_Hor_3 
                HorUpperToText_English_Hor_3 = HorUpperToText_Hor_3 
                HorLowerFromText_English_Hor_3 = HorLowerFromText_Hor_3 
                HorLowerToText_English_Hor_3 = HorLowerToText_Hor_3 
                BoundaryDistText_English_Hor_3 = proxyModelRowIndex(self.horBoundaryDistModel, BoundaryDistText_Hor_3, self.languageComboBox.currentIndex())
                BoundaryTopoText_English_Hor_3 = proxyModelRowIndex(self.horBoundaryTopoModel, BoundaryTopoText_Hor_3, self.languageComboBox.currentIndex())
                SoilColor_1_Code_Hor_3 = (SoilColorHue_1_Hor_3 + SoilColorValue_1_Hor_3 + "/"+SoilColorChroma_1_Hor_3)
                SoilColor_1_Text_English_Hor_3 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_1_Code_Hor_3, self.languageComboBox.currentIndex() + 4)
                SoilColor_2_Code_Hor_3 = (SoilColorHue_2_Hor_3 + SoilColorValue_2_Hor_3 + "/"+SoilColorChroma_2_Hor_3)
                SoilColor_2_Text_English_Hor_3 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_2_Code_Hor_3, self.languageComboBox.currentIndex() + 4)
                SoilColor_3_Code_Hor_3 = (SoilColorHue_3_Hor_3 + SoilColorValue_3_Hor_3 + "/"+SoilColorChroma_3_Hor_3)
                SoilColor_3_Text_English_Hor_3 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_3_Code_Hor_3, self.languageComboBox.currentIndex() + 4)  
                TextureClassText_English_Hor_3 = proxyModelRowIndexRegX(self.textureClassModel, TextureClassText_Hor_3, self.languageComboBox.currentIndex())
                TextureSandText_English_Hor_3 = proxyModelRowIndexRegX(self.textureSandModel, TextureSandText_Hor_3, self.languageComboBox.currentIndex())
                TextureModSizeText_English_Hor_3 = proxyModelRowIndex(self.modSizeModel, TextureModSizeText_Hor_3, self.languageComboBox.currentIndex())
                TextureModAbundanceText_English_Hor_3 = proxyModelRowIndex(self.modAbundanceModel, TextureModAbundanceText_Hor_3, self.languageComboBox.currentIndex())
                StructureShape_1_Text_English_Hor_3 = proxyModelRowIndexRegX(self.structureShapeModel, StructureShape_1_Text_Hor_3, self.languageComboBox.currentIndex())
                StructureSize_1_Text_English_Hor_3 = proxyModelRowIndexRegX(self.structureSizeModel, StructureSize_1_Text_Hor_3, self.languageComboBox.currentIndex())
                StructureGrade_1_Text_English_Hor_3 =  proxyModelRowIndex(self.structureGradeModel, StructureGrade_1_Text_Hor_3, self.languageComboBox.currentIndex())
                StructureRelationText_English_Hor_3 = proxyModelRowIndex(self.structureRelationModel, StructureRelationText_Hor_3, self.languageComboBox.currentIndex())
                StructureShape_2_Text_English_Hor_3 = proxyModelRowIndexRegX(self.structureShapeModel, StructureShape_2_Text_Hor_3, self.languageComboBox.currentIndex()) 
                StructureSize_2_Text_English_Hor_3 = proxyModelRowIndexRegX(self.structureSizeModel, StructureSize_2_Text_Hor_3, self.languageComboBox.currentIndex())
                StructureGrade_2_Text_English_Hor_3 = proxyModelRowIndex(self.structureGradeModel, StructureGrade_2_Text_Hor_3, self.languageComboBox.currentIndex())
                ConsistenceMoistText_English_Hor_3 = proxyModelRowIndexRegX(self.consistencyMoistModel, ConsistenceMoistText_Hor_3, self.languageComboBox.currentIndex())
                ConsistenceStickinessText_English_Hor_3 = proxyModelRowIndexRegX(self.consistencyStickinessModel, ConsistenceStickinessText_Hor_3, self.languageComboBox.currentIndex())
                ConsistencePlasticityText_English_Hor_3 = proxyModelRowIndexRegX(self.consistencyPlasticityModel, ConsistencePlasticityText_Hor_3, self.languageComboBox.currentIndex())
                MottleAbundance_1_Text_English_Hor_3 = proxyModelRowIndex(self.mottleAbundanceModel, MottleAbundance_1_Text_Hor_3, self.languageComboBox.currentIndex())
                MottleSize_1_Text_English_Hor_3 = proxyModelRowIndex(self.mottleSizeModel, MottleSize_1_Text_Hor_3, self.languageComboBox.currentIndex())
                MottleContrast_1_Text_English_Hor_3 = proxyModelRowIndex(self.mottleContrastModel, MottleContrast_1_Text_Hor_3, self.languageComboBox.currentIndex())
                MottleShape_1_Text_English_Hor_3 = proxyModelRowIndex(self.mottleShapeModel, MottleShape_1_Text_Hor_3, self.languageComboBox.currentIndex())
                MottleColor_1_Code_Hor_3 = (MottleHue_1_Text_Hor_3 + MottleValue_1_Text_Hor_3 + "/" +MottleChroma_1_Text_Hor_3)
                MottleColor_1_Text_English_Hor_3 = proxyModelRowIndexRegXMunsell(self.munsellModel, MottleColor_1_Code_Hor_3, self.languageComboBox.currentIndex() + 4)
                RootFineText_English_Hor_3 = proxyModelRowIndex(self.rootsFineModel, RootFineText_Hor_3, self.languageComboBox.currentIndex())
                RootMediumText_English_Hor_3 = proxyModelRowIndex(self.rootsMediumModel, RootMediumText_Hor_3, self.languageComboBox.currentIndex())
                RootCoarseText_English_Hor_3 = proxyModelRowIndex(self.rootsCoarseModel, RootCoarseText_Hor_3, self.languageComboBox.currentIndex())

                ## Horizon 4
                HorDesignDisconText_English_Hor_4 = HorDesignDisconText_Hor_4 
                HorDesignMasterText_English_Hor_4 = HorDesignMasterText_Hor_4 
                HorDesignSubText_English_Hor_4 = HorDesignSubText_Hor_4 
                HorDesignNumberText_English_Hor_4 = HorDesignNumberText_Hor_4 
                HorUpperFromText_English_Hor_4 = HorUpperFromText_Hor_4 
                HorUpperToText_English_Hor_4 = HorUpperToText_Hor_4 
                HorLowerFromText_English_Hor_4 = HorLowerFromText_Hor_4 
                HorLowerToText_English_Hor_4 = HorLowerToText_Hor_4 
                BoundaryDistText_English_Hor_4 = proxyModelRowIndex(self.horBoundaryDistModel, BoundaryDistText_Hor_4, self.languageComboBox.currentIndex())
                BoundaryTopoText_English_Hor_4 = proxyModelRowIndex(self.horBoundaryTopoModel, BoundaryTopoText_Hor_4, self.languageComboBox.currentIndex())
                SoilColor_1_Code_Hor_4 = (SoilColorHue_1_Hor_4 + SoilColorValue_1_Hor_4 + "/"+SoilColorChroma_1_Hor_4)
                SoilColor_1_Text_English_Hor_4 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_1_Code_Hor_4, self.languageComboBox.currentIndex() + 4)
                SoilColor_2_Code_Hor_4 = (SoilColorHue_2_Hor_4 + SoilColorValue_2_Hor_4 + "/"+SoilColorChroma_2_Hor_4)
                SoilColor_2_Text_English_Hor_4 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_2_Code_Hor_4, self.languageComboBox.currentIndex() + 4)
                SoilColor_3_Code_Hor_4 = (SoilColorHue_3_Hor_4 + SoilColorValue_3_Hor_4 + "/"+SoilColorChroma_3_Hor_4)
                SoilColor_3_Text_English_Hor_4 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_3_Code_Hor_4, self.languageComboBox.currentIndex() + 4)  
                TextureClassText_English_Hor_4 = proxyModelRowIndexRegX(self.textureClassModel, TextureClassText_Hor_4, self.languageComboBox.currentIndex())
                TextureSandText_English_Hor_4 = proxyModelRowIndexRegX(self.textureSandModel, TextureSandText_Hor_4, self.languageComboBox.currentIndex())
                TextureModSizeText_English_Hor_4 = proxyModelRowIndex(self.modSizeModel, TextureModSizeText_Hor_4, self.languageComboBox.currentIndex())
                TextureModAbundanceText_English_Hor_4 = proxyModelRowIndex(self.modAbundanceModel, TextureModAbundanceText_Hor_4, self.languageComboBox.currentIndex())
                StructureShape_1_Text_English_Hor_4 = proxyModelRowIndexRegX(self.structureShapeModel, StructureShape_1_Text_Hor_4, self.languageComboBox.currentIndex())
                StructureSize_1_Text_English_Hor_4 = proxyModelRowIndexRegX(self.structureSizeModel, StructureSize_1_Text_Hor_4, self.languageComboBox.currentIndex())
                StructureGrade_1_Text_English_Hor_4 =  proxyModelRowIndex(self.structureGradeModel, StructureGrade_1_Text_Hor_4, self.languageComboBox.currentIndex())
                StructureRelationText_English_Hor_4 = proxyModelRowIndex(self.structureRelationModel, StructureRelationText_Hor_4, self.languageComboBox.currentIndex())
                StructureShape_2_Text_English_Hor_4 = proxyModelRowIndexRegX(self.structureShapeModel, StructureShape_2_Text_Hor_4, self.languageComboBox.currentIndex()) 
                StructureSize_2_Text_English_Hor_4 = proxyModelRowIndexRegX(self.structureSizeModel, StructureSize_2_Text_Hor_4, self.languageComboBox.currentIndex())
                StructureGrade_2_Text_English_Hor_4 = proxyModelRowIndex(self.structureGradeModel, StructureGrade_2_Text_Hor_4, self.languageComboBox.currentIndex())
                ConsistenceMoistText_English_Hor_4 = proxyModelRowIndexRegX(self.consistencyMoistModel, ConsistenceMoistText_Hor_4, self.languageComboBox.currentIndex())
                ConsistenceStickinessText_English_Hor_4 = proxyModelRowIndexRegX(self.consistencyStickinessModel, ConsistenceStickinessText_Hor_4, self.languageComboBox.currentIndex())
                ConsistencePlasticityText_English_Hor_4 = proxyModelRowIndexRegX(self.consistencyPlasticityModel, ConsistencePlasticityText_Hor_4, self.languageComboBox.currentIndex())
                MottleAbundance_1_Text_English_Hor_4 = proxyModelRowIndex(self.mottleAbundanceModel, MottleAbundance_1_Text_Hor_4, self.languageComboBox.currentIndex())
                MottleSize_1_Text_English_Hor_4 = proxyModelRowIndex(self.mottleSizeModel, MottleSize_1_Text_Hor_4, self.languageComboBox.currentIndex())
                MottleContrast_1_Text_English_Hor_4 = proxyModelRowIndex(self.mottleContrastModel, MottleContrast_1_Text_Hor_4, self.languageComboBox.currentIndex())
                MottleShape_1_Text_English_Hor_4 = proxyModelRowIndex(self.mottleShapeModel, MottleShape_1_Text_Hor_4, self.languageComboBox.currentIndex())
                MottleColor_1_Code_Hor_4 = (MottleHue_1_Text_Hor_4 + MottleValue_1_Text_Hor_4 + "/" +MottleChroma_1_Text_Hor_4)
                MottleColor_1_Text_English_Hor_4 = proxyModelRowIndexRegXMunsell(self.munsellModel, MottleColor_1_Code_Hor_4, self.languageComboBox.currentIndex() + 4)
                RootFineText_English_Hor_4 = proxyModelRowIndex(self.rootsFineModel, RootFineText_Hor_4, self.languageComboBox.currentIndex())
                RootMediumText_English_Hor_4 = proxyModelRowIndex(self.rootsMediumModel, RootMediumText_Hor_4, self.languageComboBox.currentIndex())
                RootCoarseText_English_Hor_4 = proxyModelRowIndex(self.rootsCoarseModel, RootCoarseText_Hor_4, self.languageComboBox.currentIndex())
              
                ## Horizon 5
                HorDesignDisconText_English_Hor_5 = HorDesignDisconText_Hor_5 
                HorDesignMasterText_English_Hor_5 = HorDesignMasterText_Hor_5 
                HorDesignSubText_English_Hor_5 = HorDesignSubText_Hor_5 
                HorDesignNumberText_English_Hor_5 = HorDesignNumberText_Hor_5 
                HorUpperFromText_English_Hor_5 = HorUpperFromText_Hor_5 
                HorUpperToText_English_Hor_5 = HorUpperToText_Hor_5 
                HorLowerFromText_English_Hor_5 = HorLowerFromText_Hor_5 
                HorLowerToText_English_Hor_5 = HorLowerToText_Hor_5 
                BoundaryDistText_English_Hor_5 = proxyModelRowIndex(self.horBoundaryDistModel, BoundaryDistText_Hor_5, self.languageComboBox.currentIndex())
                BoundaryTopoText_English_Hor_5 = proxyModelRowIndex(self.horBoundaryTopoModel, BoundaryTopoText_Hor_5, self.languageComboBox.currentIndex())
                SoilColor_1_Code_Hor_5 = (SoilColorHue_1_Hor_5 + SoilColorValue_1_Hor_5 + "/"+SoilColorChroma_1_Hor_5)
                SoilColor_1_Text_English_Hor_5 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_1_Code_Hor_5, self.languageComboBox.currentIndex() + 4)
                SoilColor_2_Code_Hor_5 = (SoilColorHue_2_Hor_5 + SoilColorValue_2_Hor_5 + "/"+SoilColorChroma_2_Hor_5)
                SoilColor_2_Text_English_Hor_5 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_2_Code_Hor_5, self.languageComboBox.currentIndex() + 4)
                SoilColor_3_Code_Hor_5 = (SoilColorHue_3_Hor_5 + SoilColorValue_3_Hor_5 + "/"+SoilColorChroma_3_Hor_5)
                SoilColor_3_Text_English_Hor_5 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_3_Code_Hor_5, self.languageComboBox.currentIndex() + 4)  
                TextureClassText_English_Hor_5 = proxyModelRowIndexRegX(self.textureClassModel, TextureClassText_Hor_5, self.languageComboBox.currentIndex())
                TextureSandText_English_Hor_5 = proxyModelRowIndexRegX(self.textureSandModel, TextureSandText_Hor_5, self.languageComboBox.currentIndex())
                TextureModSizeText_English_Hor_5 = proxyModelRowIndex(self.modSizeModel, TextureModSizeText_Hor_5, self.languageComboBox.currentIndex())
                TextureModAbundanceText_English_Hor_5 = proxyModelRowIndex(self.modAbundanceModel, TextureModAbundanceText_Hor_5, self.languageComboBox.currentIndex())
                StructureShape_1_Text_English_Hor_5 = proxyModelRowIndexRegX(self.structureShapeModel, StructureShape_1_Text_Hor_5, self.languageComboBox.currentIndex())
                StructureSize_1_Text_English_Hor_5 = proxyModelRowIndexRegX(self.structureSizeModel, StructureSize_1_Text_Hor_5, self.languageComboBox.currentIndex())
                StructureGrade_1_Text_English_Hor_5 =  proxyModelRowIndex(self.structureGradeModel, StructureGrade_1_Text_Hor_5, self.languageComboBox.currentIndex())
                StructureRelationText_English_Hor_5 = proxyModelRowIndex(self.structureRelationModel, StructureRelationText_Hor_5, self.languageComboBox.currentIndex())
                StructureShape_2_Text_English_Hor_5 = proxyModelRowIndexRegX(self.structureShapeModel, StructureShape_2_Text_Hor_5, self.languageComboBox.currentIndex()) 
                StructureSize_2_Text_English_Hor_5 = proxyModelRowIndexRegX(self.structureSizeModel, StructureSize_2_Text_Hor_5, self.languageComboBox.currentIndex())
                StructureGrade_2_Text_English_Hor_5 = proxyModelRowIndex(self.structureGradeModel, StructureGrade_2_Text_Hor_5, self.languageComboBox.currentIndex())
                ConsistenceMoistText_English_Hor_5 = proxyModelRowIndexRegX(self.consistencyMoistModel, ConsistenceMoistText_Hor_5, self.languageComboBox.currentIndex())
                ConsistenceStickinessText_English_Hor_5 = proxyModelRowIndexRegX(self.consistencyStickinessModel, ConsistenceStickinessText_Hor_5, self.languageComboBox.currentIndex())
                ConsistencePlasticityText_English_Hor_5 = proxyModelRowIndexRegX(self.consistencyPlasticityModel, ConsistencePlasticityText_Hor_5, self.languageComboBox.currentIndex())
                MottleAbundance_1_Text_English_Hor_5 = proxyModelRowIndex(self.mottleAbundanceModel, MottleAbundance_1_Text_Hor_5, self.languageComboBox.currentIndex())
                MottleSize_1_Text_English_Hor_5 = proxyModelRowIndex(self.mottleSizeModel, MottleSize_1_Text_Hor_5, self.languageComboBox.currentIndex())
                MottleContrast_1_Text_English_Hor_5 = proxyModelRowIndex(self.mottleContrastModel, MottleContrast_1_Text_Hor_5, self.languageComboBox.currentIndex())
                MottleShape_1_Text_English_Hor_5 = proxyModelRowIndex(self.mottleShapeModel, MottleShape_1_Text_Hor_5, self.languageComboBox.currentIndex())
                MottleColor_1_Code_Hor_5 = (MottleHue_1_Text_Hor_5 + MottleValue_1_Text_Hor_5 + "/" +MottleChroma_1_Text_Hor_5)
                MottleColor_1_Text_English_Hor_5 = proxyModelRowIndexRegXMunsell(self.munsellModel, MottleColor_1_Code_Hor_5, self.languageComboBox.currentIndex() + 4)
                RootFineText_English_Hor_5 = proxyModelRowIndex(self.rootsFineModel, RootFineText_Hor_5, self.languageComboBox.currentIndex())
                RootMediumText_English_Hor_5 = proxyModelRowIndex(self.rootsMediumModel, RootMediumText_Hor_5, self.languageComboBox.currentIndex())
                RootCoarseText_English_Hor_5 = proxyModelRowIndex(self.rootsCoarseModel, RootCoarseText_Hor_5, self.languageComboBox.currentIndex())
              
                ## Horizon 6
                HorDesignDisconText_English_Hor_6 = HorDesignDisconText_Hor_6 
                HorDesignMasterText_English_Hor_6 = HorDesignMasterText_Hor_6 
                HorDesignSubText_English_Hor_6 = HorDesignSubText_Hor_6 
                HorDesignNumberText_English_Hor_6 = HorDesignNumberText_Hor_6 
                HorUpperFromText_English_Hor_6 = HorUpperFromText_Hor_6 
                HorUpperToText_English_Hor_6 = HorUpperToText_Hor_6 
                HorLowerFromText_English_Hor_6 = HorLowerFromText_Hor_6 
                HorLowerToText_English_Hor_6 = HorLowerToText_Hor_6 
                BoundaryDistText_English_Hor_6 = proxyModelRowIndex(self.horBoundaryDistModel, BoundaryDistText_Hor_6, self.languageComboBox.currentIndex())
                BoundaryTopoText_English_Hor_6 = proxyModelRowIndex(self.horBoundaryTopoModel, BoundaryTopoText_Hor_6, self.languageComboBox.currentIndex())
                SoilColor_1_Code_Hor_6 = (SoilColorHue_1_Hor_6 + SoilColorValue_1_Hor_6 + "/"+SoilColorChroma_1_Hor_6)
                SoilColor_1_Text_English_Hor_6 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_1_Code_Hor_6, self.languageComboBox.currentIndex() + 4)
                SoilColor_2_Code_Hor_6 = (SoilColorHue_2_Hor_6 + SoilColorValue_2_Hor_6 + "/"+SoilColorChroma_2_Hor_6)
                SoilColor_2_Text_English_Hor_6 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_2_Code_Hor_6, self.languageComboBox.currentIndex() + 4)
                SoilColor_3_Code_Hor_6 = (SoilColorHue_3_Hor_6 + SoilColorValue_3_Hor_6 + "/"+SoilColorChroma_3_Hor_6)
                SoilColor_3_Text_English_Hor_6 = proxyModelRowIndexRegXMunsell(self.munsellModel, SoilColor_3_Code_Hor_6, self.languageComboBox.currentIndex() + 4)  
                TextureClassText_English_Hor_6 = proxyModelRowIndexRegX(self.textureClassModel, TextureClassText_Hor_6, self.languageComboBox.currentIndex())
                TextureSandText_English_Hor_6 = proxyModelRowIndexRegX(self.textureSandModel, TextureSandText_Hor_6, self.languageComboBox.currentIndex())
                TextureModSizeText_English_Hor_6 = proxyModelRowIndex(self.modSizeModel, TextureModSizeText_Hor_6, self.languageComboBox.currentIndex())
                TextureModAbundanceText_English_Hor_6 = proxyModelRowIndex(self.modAbundanceModel, TextureModAbundanceText_Hor_6, self.languageComboBox.currentIndex())
                StructureShape_1_Text_English_Hor_6 = proxyModelRowIndexRegX(self.structureShapeModel, StructureShape_1_Text_Hor_6, self.languageComboBox.currentIndex())
                StructureSize_1_Text_English_Hor_6 = proxyModelRowIndexRegX(self.structureSizeModel, StructureSize_1_Text_Hor_6, self.languageComboBox.currentIndex())
                StructureGrade_1_Text_English_Hor_6 =  proxyModelRowIndex(self.structureGradeModel, StructureGrade_1_Text_Hor_6, self.languageComboBox.currentIndex())
                StructureRelationText_English_Hor_6 = proxyModelRowIndex(self.structureRelationModel, StructureRelationText_Hor_6, self.languageComboBox.currentIndex())
                StructureShape_2_Text_English_Hor_6 = proxyModelRowIndexRegX(self.structureShapeModel, StructureShape_2_Text_Hor_6, self.languageComboBox.currentIndex()) 
                StructureSize_2_Text_English_Hor_6 = proxyModelRowIndexRegX(self.structureSizeModel, StructureSize_2_Text_Hor_6, self.languageComboBox.currentIndex())
                StructureGrade_2_Text_English_Hor_6 = proxyModelRowIndex(self.structureGradeModel, StructureGrade_2_Text_Hor_6, self.languageComboBox.currentIndex())
                ConsistenceMoistText_English_Hor_6 = proxyModelRowIndexRegX(self.consistencyMoistModel, ConsistenceMoistText_Hor_6, self.languageComboBox.currentIndex())
                ConsistenceStickinessText_English_Hor_6 = proxyModelRowIndexRegX(self.consistencyStickinessModel, ConsistenceStickinessText_Hor_6, self.languageComboBox.currentIndex())
                ConsistencePlasticityText_English_Hor_6 = proxyModelRowIndexRegX(self.consistencyPlasticityModel, ConsistencePlasticityText_Hor_6, self.languageComboBox.currentIndex())
                MottleAbundance_1_Text_English_Hor_6 = proxyModelRowIndex(self.mottleAbundanceModel, MottleAbundance_1_Text_Hor_6, self.languageComboBox.currentIndex())
                MottleSize_1_Text_English_Hor_6 = proxyModelRowIndex(self.mottleSizeModel, MottleSize_1_Text_Hor_6, self.languageComboBox.currentIndex())
                MottleContrast_1_Text_English_Hor_6 = proxyModelRowIndex(self.mottleContrastModel, MottleContrast_1_Text_Hor_6, self.languageComboBox.currentIndex())
                MottleShape_1_Text_English_Hor_6 = proxyModelRowIndex(self.mottleShapeModel, MottleShape_1_Text_Hor_6, self.languageComboBox.currentIndex())
                MottleColor_1_Code_Hor_6 = (MottleHue_1_Text_Hor_6 + MottleValue_1_Text_Hor_6 + "/" +MottleChroma_1_Text_Hor_6)
                MottleColor_1_Text_English_Hor_6 = proxyModelRowIndexRegXMunsell(self.munsellModel, MottleColor_1_Code_Hor_6, self.languageComboBox.currentIndex() + 4)
                RootFineText_English_Hor_6 = proxyModelRowIndex(self.rootsFineModel, RootFineText_Hor_6, self.languageComboBox.currentIndex())
                RootMediumText_English_Hor_6 = proxyModelRowIndex(self.rootsMediumModel, RootMediumText_Hor_6, self.languageComboBox.currentIndex())
                RootCoarseText_English_Hor_6 = proxyModelRowIndex(self.rootsCoarseModel, RootCoarseText_Hor_6, self.languageComboBox.currentIndex())
              
                ### Python-Docx Implementation
                # add p paragraph
                p = document.add_paragraph("Form No./Mapping Unit")
                ### apply the font of Tahoma, size= 11, lineSpacing = 1, tabStop = 1.67, spaceBefore and spaceAfter = Pt(0)
                p.style = document.styles['Normal']
                paragraphFormat = p.paragraph_format
                paragraphFormat.space_before = Pt(0)
                paragraphFormat.space_after = Pt(0)
                paragraphFormat.line_spacing = 1
                tab_stops = paragraphFormat.tab_stops
                tab_stops.add_tab_stop(Inches(1.67))
                # add the data required 
                p.add_run("\t")
                p.add_run(': ')
                p.add_run(numberFormTextEnglish) # Number Form
                p.add_run("/")
                p.add_run(sptTextEnglish) # SPT
                p.add_run("\n")

                p.add_run("Soil Survey Staff, 2014")
                p.add_run("\t")
                p.add_run(': ')
                p.add_run(pscTextEnglish) # PSC
                p.add_run(", ")
                p.add_run(mineralTextEnglish) # Mineral
                p.add_run(", ")
                p.add_run(temperatureTextEnglish) # Temperature Regime
                p.add_run(", ")
                p.add_run(subGroupTextEnglish) # Sub Group
                p.add_run(" ")
                p.add_run(greatGroupTextEnglish) # Great Group
                p.add_run("\n")

                p.add_run("Landform")
                p.add_run("\t")
                p.add_run(': ')
                p.add_run(landformTextEnglish) # Landform
                p.add_run("\n")

                p.add_run("Parent Material")
                p.add_run("\t")
                p.add_run(': ')
                p.add_run(lithology1TextEnglish) # Lithology 1
                # Check if lithology2tex is not 'None' remove comma
                if (lithology2TextEnglish != None):
                    p.add_run(", ")
                    p.add_run(lithology2TextEnglish) # Lithology 2
                p.add_run("\n")

                p.add_run("Slope and position")
                p.add_run("\t")
                p.add_run(': ')
                p.add_run(slopePercentageTextEnglish) # Slope percentage
                p.add_run("%") 
                p.add_run(", ")
                p.add_run(slopePositionTextEnglish) # Slope position
                p.add_run("\n")

                p.add_run("Elevation")
                p.add_run("\t")
                p.add_run(': ')
                p.add_run(elevationTextEnglish) # Elevation
                p.add_run(" m")
                p.add_run(" amsl")
                p.add_run("\n")

                p.add_run("Drainage class")
                p.add_run("\t")
                p.add_run(': ')
                p.add_run(drainageClassTextEnglish) # Drainage class

                ### setting another paragraph because it have different format from above
                pLocation = document.add_paragraph()
                # apply the font of Tahoma, size= 11, lineSpacing = 1, tabStop = 1.25, spaceBefore and spaceAfter = Pt(0)
                pLocation.style = document.styles['Normal']
                pFormatLocation = pLocation.paragraph_format
                pFormatLocation.space_before = Pt(0)
                pFormatLocation.space_after = Pt(0)
                pFormatLocation.line_spacing = 1
                tabStopsLocation = pFormatLocation.tab_stops
                tabStopsLocation.add_tab_stop(Inches(1.25))

                # data required
                pLocation.add_run("Location")
                pLocation.add_run("\t")
                pLocation.add_run(":")

                ### setting another paragraph because it have different format from above

                pAdministration = document.add_paragraph()
                # apply the font of Tahoma, size= 11, lineSpacing = 1, tabStop = 1.67, spaceBefore and spaceAfter = Pt(0)
                pAdministration.style = document.styles['Normal']
                pFormatAdministration = pAdministration.paragraph_format
                pFormatAdministration.space_before = Pt(0)
                pFormatAdministration.space_after = Pt(0)
                pFormatAdministration.line_spacing = 1
                tabStopsLocation = pFormatAdministration.tab_stops
                tabStopsLocation.add_tab_stop(Inches(1.67))
                # add the data
                pAdministration.add_run(" - Administration")
                pAdministration.add_run("\t")
                pAdministration.add_run(": ")
                pAdministration.add_run(provinsiTextEnglish) # Province
                pAdministration.add_run(" Province, ")
                pAdministration.add_run(kabupatenTextEnglish) # Kabupaten/Kota
                pAdministration.add_run(" Regency, ")

                ### setting another paragraph because it have different format from above
                pKecDes = document.add_paragraph() 
                # apply the font of Tahoma, size= 11, lineSpacing = 1, tabStop = 1.77, spaceBefore and spaceAfter = Pt(0)
                pKecDes.style = document.styles['Normal']
                pFormatKecDes = pKecDes.paragraph_format
                pFormatKecDes.space_before = Pt(0)
                pFormatKecDes.space_after = Pt(0)
                pKecDes.paragraph_format.first_line_indent = Inches(1.77)
                pKecDes.add_run(kecamatanTextEnglish) # Kecamatan
                pKecDes.add_run(" Sub-Regency")
                pKecDes.add_run(", ")
                pKecDes.add_run(desaTextEnglish) # Desa/Kompartemen

                ## setting another paragraph because it have different format from above
                pCoordinate = document.add_paragraph()
                # apply the font of Tahoma, size= 11, lineSpacing = 1, tabStop = 1.67, spaceBefore and spaceAfter = Pt(0)
                pCoordinate.style = document.styles['Normal']
                pFormatCoordinate = pCoordinate.paragraph_format
                pFormatCoordinate.space_before = Pt(0)
                pFormatCoordinate.space_after = Pt(0)
                pFormatCoordinate.line_spacing = 1
                tabStopsCoordinate = pFormatCoordinate.tab_stops
                tabStopsCoordinate.add_tab_stop(Inches(1.67))
                # add the data
                pCoordinate.add_run(" - Coordinate")
                pCoordinate.add_run("\t")
                pCoordinate.add_run(": ")
                pCoordinate.add_run("X: ")
                pCoordinate.add_run(xEastTextEnglish) # X-East
                pCoordinate.add_run(" (E), Y: ")
                pCoordinate.add_run(yNorthTextEnglish) # Y-North
                pCoordinate.add_run(" (N) - Zone: ")
                pCoordinate.add_run(utmZone1TextEnglish+utmZone2TextEnglish) # UTM Zone
                ## setting another paragraph because it have different format from above
                pSiteNumber = document.add_paragraph()
                # apply the font of Tahoma, size= 11, lineSpacing = 1, tabStop = 1.67, spaceBefore and spaceAfter = Pt(0)
                pSiteNumber.style = document.styles['Normal']
                pFormatSiteNumber = pSiteNumber.paragraph_format
                pFormatSiteNumber.space_before = Pt(0)
                pFormatSiteNumber.space_after = Pt(0)
                pFormatSiteNumber.line_spacing = 1
                tabStopsSiteNumber = pFormatSiteNumber.tab_stops
                tabStopsSiteNumber.add_tab_stop(Inches(1.67))
                # add the data
                pSiteNumber.add_run(" - Site Number")
                pSiteNumber.add_run("\t")
                pSiteNumber.add_run(": ")
                pSiteNumber.add_run(dateTextEnglish) # Date
                pSiteNumber.add_run(" ") 
                pSiteNumber.add_run(initialNameTextEnglish) # Initial Name
                pSiteNumber.add_run("<")
                pSiteNumber.add_run(observationNumberTextEnglish) # Observation number
                pSiteNumber.add_run(">")
                pSiteNumber.add_run("; ")
                pSiteNumber.add_run(kindObservationTextEnglish) # Type of Observations
                pSiteNumber.add_run("\n")
                
                # # add picture, but the feature of wrapped text is not yet implemented
                # document.add_picture("I:\_Python\PyQT\SiteandHorizon\\New\picture\FM066.JPG", width = Inches(1.9))
               

                ### HORIZON I implementation ###

                # setting another paragraph because it have different format from above
                horizon1Paragraph = document.add_paragraph()
                # apply the font of Tahoma, size= 11, leftIndent = 2.45, first_line_indent = - 5 from left_indent,
                # align with justify_low, spaceBefore and spaceAfter = Pt(0)
                horizon1Paragraph.style = document.styles['Normal']
                horizon1ParagraphFormat = horizon1Paragraph.paragraph_format
                horizon1ParagraphFormat.line_spacing = 1
                horizon1ParagraphFormat.space_before = Pt(6)
                horizon1ParagraphFormat.space_after = Pt(0)
                horizon1Paragraph.paragraph_format.left_indent = Inches (2.5)
                horizon1Paragraph.paragraph_format.first_line_indent = Inches(-0.5)
                horizon1ParagraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


                ### add the data ###
                # Horison 
                if (horizonRow != -1):
                    if (HorDesignMasterText_English_Hor_1 != ""):
                        horizon1Paragraph.add_run(HorDesignDisconText_English_Hor_1 + HorDesignMasterText_English_Hor_1 + HorDesignSubText_English_Hor_1 + HorDesignNumberText_English_Hor_1)# horison design desig+master+sub+number
                        horizon1Paragraph.add_run("\t")
                    
                    # Depth
                    if (HorUpperToText_English_Hor_1 != ""):
                        horizon1Paragraph.add_run(HorUpperToText_English_Hor_1)# Upper limit to
                        horizon1Paragraph.add_run("-")
                    if (HorLowerToText_English_Hor_1 != ""):
                        horizon1Paragraph.add_run(HorLowerToText_English_Hor_1)# Lower limit to
                        horizon1Paragraph.add_run(" cm")
                        horizon1Paragraph.add_run("; ")

                # Soil Color 1
                if (SoilColor_1_Code_Hor_1 != "/" and SoilColor_1_Text_English_Hor_1 != None):
                    horizon1Paragraph.add_run(SoilColor_1_Text_English_Hor_1) # munsel color text english 
                    horizon1Paragraph.add_run(" (") 
                    horizon1Paragraph.add_run(SoilColor_1_Code_Hor_1) # Munsel Hue + Value + / + Chroma
                    horizon1Paragraph.add_run("); ") # 

                # Soil Color 2
                if (SoilColor_2_Code_Hor_1 != "/" and SoilColor_2_Text_English_Hor_1 != None):
                    horizon1Paragraph.add_run(SoilColor_2_Text_English_Hor_1) # munsel color text english
                    horizon1Paragraph.add_run(" (") 
                    horizon1Paragraph.add_run(SoilColor_2_Code_Hor_1) # Munsel Hue + Value + / + Chroma
                    horizon1Paragraph.add_run("); ") # 

                # Soil Color 3
                if (SoilColor_3_Code_Hor_1 != "/" and SoilColor_3_Text_English_Hor_1 != None):
                    horizon1Paragraph.add_run(SoilColor_3_Text_English_Hor_1) # munsel color text english
                    horizon1Paragraph.add_run(" (") 
                    horizon1Paragraph.add_run(SoilColor_3_Code_Hor_1) # Munsel Hue + Value + / + Chroma
                    horizon1Paragraph.add_run("); ") # 

                # Texture
                if (TextureClassText_English_Hor_1 != None):
                    horizon1Paragraph.add_run(TextureClassText_English_Hor_1) # texture class
                    horizon1Paragraph.add_run("; ")
                
                # Texture Sand
                if (TextureSandText_English_Hor_1 != None):
                    horizon1Paragraph.add_run(TextureSandText_English_Hor_1) # Texture sand
                    horizon1Paragraph.add_run(" sand; ")

                
                # Texture Modifier Size
                if(TextureModSizeText_English_Hor_1 != None):
                    horizon1Paragraph.add_run(TextureModSizeText_English_Hor_1)# Texture modifier size
                    horizon1Paragraph.add_run(", ")
                
                # Texture Modifier Abundance
                if(TextureModAbundanceText_English_Hor_1 != None):
                    horizon1Paragraph.add_run(TextureModAbundanceText_English_Hor_1)# Texture modifier size
                    horizon1Paragraph.add_run(" modifier; ")
                

                # Structure 1
                if (StructureShape_1_Text_English_Hor_1 != None):
                    horizon1Paragraph.add_run(StructureGrade_1_Text_English_Hor_1) # structure grade
                    horizon1Paragraph.add_run(", ")
                    horizon1Paragraph.add_run(StructureSize_1_Text_English_Hor_1) # structure size
                    horizon1Paragraph.add_run(", ")
                    horizon1Paragraph.add_run(StructureShape_1_Text_English_Hor_1) # structure shape
                    horizon1Paragraph.add_run(" structure; ")

                # Structure Relation
                if (StructureRelationText_English_Hor_1 != None):
                    horizon1Paragraph.add_run(StructureRelationText_English_Hor_1) # Structure Relation   

                # Structure 2
                if (StructureShape_2_Text_English_Hor_1 != None):
                    horizon1Paragraph.add_run(StructureGrade_2_Text_English_Hor_1) # structure grade
                    horizon1Paragraph.add_run(", ")
                    horizon1Paragraph.add_run(StructureSize_2_Text_English_Hor_1) # structure size
                    horizon1Paragraph.add_run(", ")
                    horizon1Paragraph.add_run(StructureShape_2_Text_English_Hor_1) # structure shape
                    horizon1Paragraph.add_run(" structure; ")

                # Consistence
                if (ConsistenceMoistText_English_Hor_1 != None):
                    horizon1Paragraph.add_run(ConsistenceMoistText_English_Hor_1) # consistency moist
                    horizon1Paragraph.add_run(" (moist), ")
                    horizon1Paragraph.add_run(ConsistenceStickinessText_English_Hor_1) # consistency stickiness
                    horizon1Paragraph.add_run(" and ")
                    horizon1Paragraph.add_run(ConsistencePlasticityText_English_Hor_1) # consitency plasticity
                    horizon1Paragraph.add_run(" (wet) consistency; ")
                
                # Roots
                if (RootFineText_English_Hor_1 != None):
                    horizon1Paragraph.add_run(RootFineText_English_Hor_1) # root fine
                    horizon1Paragraph.add_run(" fine")
                    if (RootMediumText_English_Hor_1 != None):
                        horizon1Paragraph.add_run(", ")
                        horizon1Paragraph.add_run(RootMediumText_English_Hor_1) # root medium
                        horizon1Paragraph.add_run(" medium")
                    if (RootCoarseText_English_Hor_1 != None):
                        horizon1Paragraph.add_run(", ")
                        horizon1Paragraph.add_run(RootCoarseText_English_Hor_1) # root coarse
                        horizon1Paragraph.add_run(" coarse")

                    horizon1Paragraph.add_run(" roots; ")

                # Mottles
                if (MottleColor_1_Text_English_Hor_1 != None):
                    if (MottleAbundance_1_Text_English_Hor_1 != None):
                        horizon1Paragraph.add_run(MottleAbundance_1_Text_English_Hor_1) # Mottle Abundance
                        horizon1Paragraph.add_run(", ")
                    if (MottleSize_1_Text_English_Hor_1 != None):
                        horizon1Paragraph.add_run(MottleSize_1_Text_English_Hor_1) # Mottle Size
                        horizon1Paragraph.add_run(", ")
                    if (MottleContrast_1_Text_English_Hor_1 != None):
                        horizon1Paragraph.add_run(MottleContrast_1_Text_English_Hor_1) # Mottle Contrast
                        horizon1Paragraph.add_run(", ")
                    if (MottleShape_1_Text_English_Hor_1 != None):
                        horizon1Paragraph.add_run(MottleShape_1_Text_English_Hor_1) # Mottle Shape
                        horizon1Paragraph.add_run(", ")
                    horizon1Paragraph.add_run(MottleColor_1_Text_English_Hor_1) # Mottle Color English
                    horizon1Paragraph.add_run(" (")
                    horizon1Paragraph.add_run(MottleColor_1_Code_Hor_1) # Mottle Color Code (Hue+Value+/+Chroma)
                    horizon1Paragraph.add_run(")")
                    horizon1Paragraph.add_run(" mottles; ")


                # Horizon Boundary
                if (BoundaryDistText_English_Hor_1 != None):    
                    horizon1Paragraph.add_run(BoundaryDistText_English_Hor_1) # horizon boundary dist
                    horizon1Paragraph.add_run(" ")
                    horizon1Paragraph.add_run(BoundaryTopoText_English_Hor_1) # horizon boundary topo
                    horizon1Paragraph.add_run(" boundary.")

                ### HORIZON 2 implementation ###

                # setting another paragraph because it have different format from above
                horizon2Paragraph = document.add_paragraph()
                # apply the font of Tahoma, size= 11, leftIndent = 2.45, first_line_indent = - 5 from left_indent,
                # align with justify_low, spaceBefore and spaceAfter = Pt(0)
                horizon2Paragraph.style = document.styles['Normal']
                horizon2ParagraphFormat = horizon2Paragraph.paragraph_format
                horizon2ParagraphFormat.line_spacing = 1
                horizon2ParagraphFormat.space_before = Pt(6)
                horizon2ParagraphFormat.space_after = Pt(0)
                horizon2Paragraph.paragraph_format.left_indent = Inches (2.5)
                horizon2Paragraph.paragraph_format.first_line_indent = Inches(-0.5)
                horizon2ParagraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Horison
                if (horizonRow != -1): 
                    if (HorDesignMasterText_English_Hor_2 != ""):
                        horizon2Paragraph.add_run(HorDesignDisconText_English_Hor_2 + HorDesignMasterText_English_Hor_2 + HorDesignSubText_English_Hor_2 + HorDesignNumberText_English_Hor_2)# horison design desig+master+sub+number
                        horizon2Paragraph.add_run("\t")
                    
                    # Depth
                    if (HorUpperToText_English_Hor_2 != ""):
                        horizon2Paragraph.add_run(HorUpperToText_English_Hor_2)# Upper limit to
                        horizon2Paragraph.add_run("-")
                    if (HorLowerToText_English_Hor_2 != ""):
                        horizon2Paragraph.add_run(HorLowerToText_English_Hor_2)# Lower limit to
                        horizon2Paragraph.add_run(" cm")
                        horizon2Paragraph.add_run("; ")

                # Soil Color 1
                if (SoilColor_1_Code_Hor_2 != "/" and SoilColor_1_Text_English_Hor_2 != None):
                    horizon2Paragraph.add_run(SoilColor_1_Text_English_Hor_2) # munsel color text english 
                    horizon2Paragraph.add_run(" (") 
                    horizon2Paragraph.add_run(SoilColor_1_Code_Hor_2) # Munsel Hue + Value + / + Chroma
                    horizon2Paragraph.add_run("); ") # 

                # Soil Color 2
                if (SoilColor_2_Code_Hor_2 != "/" and SoilColor_2_Text_English_Hor_2 != None):
                    horizon2Paragraph.add_run(SoilColor_2_Text_English_Hor_2) # munsel color text english
                    horizon2Paragraph.add_run(" (") 
                    horizon2Paragraph.add_run(SoilColor_2_Code_Hor_2) # Munsel Hue + Value + / + Chroma
                    horizon2Paragraph.add_run("); ") # 

                # Soil Color 3
                if (SoilColor_3_Code_Hor_2 != "/" and SoilColor_3_Text_English_Hor_2 != None):
                    horizon2Paragraph.add_run(SoilColor_3_Text_English_Hor_2) # munsel color text english
                    horizon2Paragraph.add_run(" (") 
                    horizon2Paragraph.add_run(SoilColor_3_Code_Hor_2) # Munsel Hue + Value + / + Chroma
                    horizon2Paragraph.add_run("); ") # 

                # Texture
                if (TextureClassText_English_Hor_2 != None):
                    horizon2Paragraph.add_run(TextureClassText_English_Hor_2) # texture class
                    horizon2Paragraph.add_run("; ")
                
                # Texture Sand
                if (TextureSandText_English_Hor_2 != None):
                    horizon2Paragraph.add_run(TextureSandText_English_Hor_2) # Texture sand
                    horizon2Paragraph.add_run(" sand; ")

                
                # Texture Modifier Size
                if(TextureModSizeText_English_Hor_2 != None):
                    horizon2Paragraph.add_run(TextureModSizeText_English_Hor_2)# Texture modifier size
                    horizon2Paragraph.add_run(", ")
                
                # Texture Modifier Abundance
                if(TextureModAbundanceText_English_Hor_2 != None):
                    horizon2Paragraph.add_run(TextureModAbundanceText_English_Hor_2)# Texture modifier size
                    horizon2Paragraph.add_run(" modifier; ")

                # Structure 1
                if (StructureShape_1_Text_English_Hor_2 != None):
                    horizon2Paragraph.add_run(StructureGrade_1_Text_English_Hor_2) # structure grade
                    horizon2Paragraph.add_run(", ")
                    horizon2Paragraph.add_run(StructureSize_1_Text_English_Hor_2) # structure size
                    horizon2Paragraph.add_run(", ")
                    horizon2Paragraph.add_run(StructureShape_1_Text_English_Hor_2) # structure shape
                    horizon2Paragraph.add_run(" structure; ")

                # Structure Relation
                if (StructureRelationText_English_Hor_2 != None):
                    horizon2Paragraph.add_run(StructureRelationText_English_Hor_2) # Structure Relation   

                # Structure 2
                if (StructureShape_2_Text_English_Hor_2 != None):
                    horizon2Paragraph.add_run(StructureGrade_2_Text_English_Hor_2) # structure grade
                    horizon2Paragraph.add_run(", ")
                    horizon2Paragraph.add_run(StructureSize_2_Text_English_Hor_2) # structure size
                    horizon2Paragraph.add_run(", ")
                    horizon2Paragraph.add_run(StructureShape_2_Text_English_Hor_2) # structure shape
                    horizon2Paragraph.add_run(" structure; ")

                # Consistence
                if (ConsistenceMoistText_English_Hor_2 != None):
                    horizon2Paragraph.add_run(ConsistenceMoistText_English_Hor_2) # consistency moist
                    horizon2Paragraph.add_run(" (moist), ")
                    horizon2Paragraph.add_run(ConsistenceStickinessText_English_Hor_2) # consistency stickiness
                    horizon2Paragraph.add_run(" and ")
                    horizon2Paragraph.add_run(ConsistencePlasticityText_English_Hor_2) # consitency plasticity
                    horizon2Paragraph.add_run(" (wet) consistency; ")
                
                # Roots
                if (RootFineText_English_Hor_2 != None):
                    horizon2Paragraph.add_run(RootFineText_English_Hor_2) # root fine
                    horizon2Paragraph.add_run(" fine")
                    if (RootMediumText_English_Hor_2 != None):
                        horizon2Paragraph.add_run(", ")
                        horizon2Paragraph.add_run(RootMediumText_English_Hor_2) # root medium
                        horizon2Paragraph.add_run(" medium")
                    if (RootCoarseText_English_Hor_2 != None):
                        horizon2Paragraph.add_run(", ")
                        horizon2Paragraph.add_run(RootCoarseText_English_Hor_2) # root coarse
                        horizon2Paragraph.add_run(" coarse")

                    horizon2Paragraph.add_run(" roots; ")

                # Mottles
                if (MottleColor_1_Text_English_Hor_2 != None):
                    if (MottleAbundance_1_Text_English_Hor_2 != None):
                        horizon2Paragraph.add_run(MottleAbundance_1_Text_English_Hor_2) # Mottle Abundance
                        horizon2Paragraph.add_run(", ")
                    if (MottleSize_1_Text_English_Hor_2 != None):
                        horizon2Paragraph.add_run(MottleSize_1_Text_English_Hor_2) # Mottle Size
                        horizon2Paragraph.add_run(", ")
                    if (MottleContrast_1_Text_English_Hor_2 != None):
                        horizon2Paragraph.add_run(MottleContrast_1_Text_English_Hor_2) # Mottle Contrast
                        horizon2Paragraph.add_run(", ")
                    if (MottleShape_1_Text_English_Hor_2 != None):
                        horizon2Paragraph.add_run(MottleShape_1_Text_English_Hor_2) # Mottle Shape
                        horizon2Paragraph.add_run(", ")
                    horizon2Paragraph.add_run(MottleColor_1_Text_English_Hor_2) # Mottle Color English
                    horizon2Paragraph.add_run(" (")
                    horizon2Paragraph.add_run(MottleColor_1_Code_Hor_2) # Mottle Color Code (Hue+Value+/+Chroma)
                    horizon2Paragraph.add_run(")")
                    horizon2Paragraph.add_run(" mottles; ")


                # Horizon Boundary
                if (BoundaryDistText_English_Hor_2 != None):    
                    horizon2Paragraph.add_run(BoundaryDistText_English_Hor_2) # horizon boundary dist
                    horizon2Paragraph.add_run(" ")
                    horizon2Paragraph.add_run(BoundaryTopoText_English_Hor_2) # horizon boundary topo
                    horizon2Paragraph.add_run(" boundary.")

                
                
                ### HORIZON 3 implementation ###

                # setting another paragraph because it have different format from above
                horizon3Paragraph = document.add_paragraph()
                # apply the font of Tahoma, size= 11, leftIndent = 2.45, first_line_indent = - 5 from left_indent,
                # align with justify_low, spaceBefore and spaceAfter = Pt(0)
                horizon3Paragraph.style = document.styles['Normal']
                horizon3ParagraphFormat = horizon3Paragraph.paragraph_format
                horizon3ParagraphFormat.line_spacing = 1
                horizon3ParagraphFormat.space_before = Pt(6)
                horizon3ParagraphFormat.space_after = Pt(0)
                horizon3Paragraph.paragraph_format.left_indent = Inches (2.5)
                horizon3Paragraph.paragraph_format.first_line_indent = Inches(-0.5)
                horizon3ParagraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Horison 
                if (horizonRow != -1):
                    if (HorDesignMasterText_English_Hor_3 != ""):
                        horizon3Paragraph.add_run(HorDesignDisconText_English_Hor_3 + HorDesignMasterText_English_Hor_3 + HorDesignSubText_English_Hor_3 + HorDesignNumberText_English_Hor_3)# horison design desig+master+sub+number
                        horizon3Paragraph.add_run("\t")
                    
                    # Depth
                    if (HorUpperToText_English_Hor_3 != ""):
                        horizon3Paragraph.add_run(HorUpperToText_English_Hor_3)# Upper limit to
                        horizon3Paragraph.add_run("-")
                    if (HorLowerToText_English_Hor_3 != ""):
                        horizon3Paragraph.add_run(HorLowerToText_English_Hor_3)# Lower limit to
                        horizon3Paragraph.add_run(" cm")
                        horizon3Paragraph.add_run("; ")

                # Soil Color 1
                if (SoilColor_1_Code_Hor_3 != "/" and SoilColor_1_Text_English_Hor_3 != None):
                    horizon3Paragraph.add_run(SoilColor_1_Text_English_Hor_3) # munsel color text english 
                    horizon3Paragraph.add_run(" (") 
                    horizon3Paragraph.add_run(SoilColor_1_Code_Hor_3) # Munsel Hue + Value + / + Chroma
                    horizon3Paragraph.add_run("); ") # 

                # Soil Color 2
                if (SoilColor_2_Code_Hor_3 != "/" and SoilColor_2_Text_English_Hor_3 != None):
                    horizon3Paragraph.add_run(SoilColor_2_Text_English_Hor_3) # munsel color text english
                    horizon3Paragraph.add_run(" (") 
                    horizon3Paragraph.add_run(SoilColor_2_Code_Hor_3) # Munsel Hue + Value + / + Chroma
                    horizon3Paragraph.add_run("); ") # 

                # Soil Color 3
                if (SoilColor_3_Code_Hor_3 != "/" and SoilColor_3_Text_English_Hor_3 != None):
                    horizon3Paragraph.add_run(SoilColor_3_Text_English_Hor_3) # munsel color text english
                    horizon3Paragraph.add_run(" (") 
                    horizon3Paragraph.add_run(SoilColor_3_Code_Hor_3) # Munsel Hue + Value + / + Chroma
                    horizon3Paragraph.add_run("); ") # 

                # Texture
                if (TextureClassText_English_Hor_3 != None):
                    horizon3Paragraph.add_run(TextureClassText_English_Hor_3) # texture class
                    horizon3Paragraph.add_run("; ")

                # Texture Sand
                if (TextureSandText_English_Hor_3 != None):
                    horizon3Paragraph.add_run(TextureSandText_English_Hor_3) # Texture sand
                    horizon3Paragraph.add_run(" sand; ")

                
                # Texture Modifier Size
                if(TextureModSizeText_English_Hor_3 != None):
                    horizon3Paragraph.add_run(TextureModSizeText_English_Hor_3)# Texture modifier size
                    horizon3Paragraph.add_run(", ")
                
                # Texture Modifier Abundance
                if(TextureModAbundanceText_English_Hor_3 != None):
                    horizon3Paragraph.add_run(TextureModAbundanceText_English_Hor_3)# Texture modifier size
                    horizon3Paragraph.add_run(" modifier; ")

                # Structure 1
                if (StructureShape_1_Text_English_Hor_3 != None):
                    horizon3Paragraph.add_run(StructureGrade_1_Text_English_Hor_3) # structure grade
                    horizon3Paragraph.add_run(", ")
                    horizon3Paragraph.add_run(StructureSize_1_Text_English_Hor_3) # structure size
                    horizon3Paragraph.add_run(", ")
                    horizon3Paragraph.add_run(StructureShape_1_Text_English_Hor_3) # structure shape
                    horizon3Paragraph.add_run(" structure; ")

                # Structure Relation
                if (StructureRelationText_English_Hor_3 != None):
                    horizon3Paragraph.add_run(StructureRelationText_English_Hor_3) # Structure Relation   

                # Structure 2
                if (StructureShape_2_Text_English_Hor_3 != None):
                    horizon3Paragraph.add_run(StructureGrade_2_Text_English_Hor_3) # structure grade
                    horizon3Paragraph.add_run(", ")
                    horizon3Paragraph.add_run(StructureSize_2_Text_English_Hor_3) # structure size
                    horizon3Paragraph.add_run(", ")
                    horizon3Paragraph.add_run(StructureShape_2_Text_English_Hor_3) # structure shape
                    horizon3Paragraph.add_run(" structure; ")

                # Consistence
                if (ConsistenceMoistText_English_Hor_3 != None):
                    horizon3Paragraph.add_run(ConsistenceMoistText_English_Hor_3) # consistency moist
                    horizon3Paragraph.add_run(" (moist), ")
                    horizon3Paragraph.add_run(ConsistenceStickinessText_English_Hor_3) # consistency stickiness
                    horizon3Paragraph.add_run(" and ")
                    horizon3Paragraph.add_run(ConsistencePlasticityText_English_Hor_3) # consitency plasticity
                    horizon3Paragraph.add_run(" (wet) consistency; ")
                
                # Roots
                if (RootFineText_English_Hor_3 != None):
                    horizon3Paragraph.add_run(RootFineText_English_Hor_3) # root fine
                    horizon3Paragraph.add_run(" fine")
                    if (RootMediumText_English_Hor_3 != None):
                        horizon3Paragraph.add_run(", ")
                        horizon3Paragraph.add_run(RootMediumText_English_Hor_3) # root medium
                        horizon3Paragraph.add_run(" medium")
                    if (RootCoarseText_English_Hor_3 != None):
                        horizon3Paragraph.add_run(", ")
                        horizon3Paragraph.add_run(RootCoarseText_English_Hor_3) # root coarse
                        horizon3Paragraph.add_run(" coarse")

                    horizon3Paragraph.add_run(" roots; ")

                # Mottles
                if (MottleColor_1_Text_English_Hor_3 != None):
                    if (MottleAbundance_1_Text_English_Hor_3 != None):
                        horizon3Paragraph.add_run(MottleAbundance_1_Text_English_Hor_3) # Mottle Abundance
                        horizon3Paragraph.add_run(", ")
                    if (MottleSize_1_Text_English_Hor_3 != None):
                        horizon3Paragraph.add_run(MottleSize_1_Text_English_Hor_3) # Mottle Size
                        horizon3Paragraph.add_run(", ")
                    if (MottleContrast_1_Text_English_Hor_3 != None):
                        horizon3Paragraph.add_run(MottleContrast_1_Text_English_Hor_3) # Mottle Contrast
                        horizon3Paragraph.add_run(", ")
                    if (MottleShape_1_Text_English_Hor_3 != None):
                        horizon3Paragraph.add_run(MottleShape_1_Text_English_Hor_3) # Mottle Shape
                        horizon3Paragraph.add_run(", ")
                    horizon3Paragraph.add_run(MottleColor_1_Text_English_Hor_3) # Mottle Color English
                    horizon3Paragraph.add_run(" (")
                    horizon3Paragraph.add_run(MottleColor_1_Code_Hor_3) # Mottle Color Code (Hue+Value+/+Chroma)
                    horizon3Paragraph.add_run(")")
                    horizon3Paragraph.add_run(" mottles; ")


                # Horizon Boundary
                if (BoundaryDistText_English_Hor_3 != None):    
                    horizon3Paragraph.add_run(BoundaryDistText_English_Hor_3) # horizon boundary dist
                    horizon3Paragraph.add_run(" ")
                    horizon3Paragraph.add_run(BoundaryTopoText_English_Hor_3) # horizon boundary topo
                    horizon3Paragraph.add_run(" boundary.")

                
                
                ### HORIZON 4 implementation ###

                # setting another paragraph because it have different format from above
                horizon4Paragraph = document.add_paragraph()
                # apply the font of Tahoma, size= 11, leftIndent = 2.45, first_line_indent = - 5 from left_indent,
                # align with justify_low, spaceBefore and spaceAfter = Pt(0)
                horizon4Paragraph.style = document.styles['Normal']
                horizon4ParagraphFormat = horizon4Paragraph.paragraph_format
                horizon4ParagraphFormat.line_spacing = 1
                horizon4ParagraphFormat.space_before = Pt(6)
                horizon4ParagraphFormat.space_after = Pt(0)
                horizon4Paragraph.paragraph_format.left_indent = Inches (2.5)
                horizon4Paragraph.paragraph_format.first_line_indent = Inches(-0.5)
                horizon4ParagraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Horison 
                if (horizonRow != -1):
                    if (HorDesignMasterText_English_Hor_4 != ""):
                        horizon4Paragraph.add_run(HorDesignDisconText_English_Hor_4 + HorDesignMasterText_English_Hor_4 + HorDesignSubText_English_Hor_4 + HorDesignNumberText_English_Hor_4)# horison design desig+master+sub+number
                        horizon4Paragraph.add_run("\t")
                    
                    # Depth
                    if (HorUpperToText_English_Hor_4 != ""):
                        horizon4Paragraph.add_run(HorUpperToText_English_Hor_4)# Upper limit to
                        horizon4Paragraph.add_run("-")
                    if (HorLowerToText_English_Hor_4 != ""):
                        horizon4Paragraph.add_run(HorLowerToText_English_Hor_4)# Lower limit to
                        horizon4Paragraph.add_run(" cm")
                        horizon4Paragraph.add_run("; ")

                # Soil Color 1
                if (SoilColor_1_Code_Hor_4 != "/" and SoilColor_1_Text_English_Hor_4 != None):
                    horizon4Paragraph.add_run(SoilColor_1_Text_English_Hor_4) # munsel color text english 
                    horizon4Paragraph.add_run(" (") 
                    horizon4Paragraph.add_run(SoilColor_1_Code_Hor_4) # Munsel Hue + Value + / + Chroma
                    horizon4Paragraph.add_run("); ") # 

                # Soil Color 2
                if (SoilColor_2_Code_Hor_4 != "/" and SoilColor_2_Text_English_Hor_4 != None):
                    horizon4Paragraph.add_run(SoilColor_2_Text_English_Hor_4) # munsel color text english
                    horizon4Paragraph.add_run(" (") 
                    horizon4Paragraph.add_run(SoilColor_2_Code_Hor_4) # Munsel Hue + Value + / + Chroma
                    horizon4Paragraph.add_run("); ") # 

                # Soil Color 3
                if (SoilColor_3_Code_Hor_4 != "/" and SoilColor_3_Text_English_Hor_4 != None):
                    horizon4Paragraph.add_run(SoilColor_3_Text_English_Hor_4) # munsel color text english
                    horizon4Paragraph.add_run(" (") 
                    horizon4Paragraph.add_run(SoilColor_3_Code_Hor_4) # Munsel Hue + Value + / + Chroma
                    horizon4Paragraph.add_run("); ") # 

                # Texture
                if (TextureClassText_English_Hor_4 != None):
                    horizon4Paragraph.add_run(TextureClassText_English_Hor_4) # texture class
                    horizon4Paragraph.add_run("; ")

                # Texture Sand
                if (TextureSandText_English_Hor_4 != None):
                    horizon4Paragraph.add_run(TextureSandText_English_Hor_4) # Texture sand
                    horizon4Paragraph.add_run(" sand; ")

                
                # Texture Modifier Size
                if(TextureModSizeText_English_Hor_4 != None):
                    horizon4Paragraph.add_run(TextureModSizeText_English_Hor_4)# Texture modifier size
                    horizon4Paragraph.add_run(", ")
                
                # Texture Modifier Abundance
                if(TextureModAbundanceText_English_Hor_4 != None):
                    horizon4Paragraph.add_run(TextureModAbundanceText_English_Hor_4)# Texture modifier size
                    horizon4Paragraph.add_run(" modifier; ")

                # Structure 1
                if (StructureShape_1_Text_English_Hor_4 != None):
                    horizon4Paragraph.add_run(StructureGrade_1_Text_English_Hor_4) # structure grade
                    horizon4Paragraph.add_run(", ")
                    horizon4Paragraph.add_run(StructureSize_1_Text_English_Hor_4) # structure size
                    horizon4Paragraph.add_run(", ")
                    horizon4Paragraph.add_run(StructureShape_1_Text_English_Hor_4) # structure shape
                    horizon4Paragraph.add_run(" structure; ")

                # Structure Relation
                if (StructureRelationText_English_Hor_4 != None):
                    horizon4Paragraph.add_run(StructureRelationText_English_Hor_4) # Structure Relation   

                # Structure 2
                if (StructureShape_2_Text_English_Hor_4 != None):
                    horizon4Paragraph.add_run(StructureGrade_2_Text_English_Hor_4) # structure grade
                    horizon4Paragraph.add_run(", ")
                    horizon4Paragraph.add_run(StructureSize_2_Text_English_Hor_4) # structure size
                    horizon4Paragraph.add_run(", ")
                    horizon4Paragraph.add_run(StructureShape_2_Text_English_Hor_4) # structure shape
                    horizon4Paragraph.add_run(" structure; ")

                # Consistence
                if (ConsistenceMoistText_English_Hor_4 != None):
                    horizon4Paragraph.add_run(ConsistenceMoistText_English_Hor_4) # consistency moist
                    horizon4Paragraph.add_run(" (moist), ")
                    horizon4Paragraph.add_run(ConsistenceStickinessText_English_Hor_4) # consistency stickiness
                    horizon4Paragraph.add_run(" and ")
                    horizon4Paragraph.add_run(ConsistencePlasticityText_English_Hor_4) # consitency plasticity
                    horizon4Paragraph.add_run(" (wet) consistency; ")
                
                # Roots
                if (RootFineText_English_Hor_4 != None):
                    horizon4Paragraph.add_run(RootFineText_English_Hor_4) # root fine
                    horizon4Paragraph.add_run(" fine")
                    if (RootMediumText_English_Hor_4 != None):
                        horizon4Paragraph.add_run(", ")
                        horizon4Paragraph.add_run(RootMediumText_English_Hor_4) # root medium
                        horizon4Paragraph.add_run(" medium")
                    if (RootCoarseText_English_Hor_4 != None):
                        horizon4Paragraph.add_run(", ")
                        horizon4Paragraph.add_run(RootCoarseText_English_Hor_4) # root coarse
                        horizon4Paragraph.add_run(" coarse")

                    horizon4Paragraph.add_run(" roots; ")

                # Mottles
                if (MottleColor_1_Text_English_Hor_4 != None):
                    if (MottleAbundance_1_Text_English_Hor_4 != None):
                        horizon4Paragraph.add_run(MottleAbundance_1_Text_English_Hor_4) # Mottle Abundance
                        horizon4Paragraph.add_run(", ")
                    if (MottleSize_1_Text_English_Hor_4 != None):
                        horizon4Paragraph.add_run(MottleSize_1_Text_English_Hor_4) # Mottle Size
                        horizon4Paragraph.add_run(", ")
                    if (MottleContrast_1_Text_English_Hor_4 != None):
                        horizon4Paragraph.add_run(MottleContrast_1_Text_English_Hor_4) # Mottle Contrast
                        horizon4Paragraph.add_run(", ")
                    if (MottleShape_1_Text_English_Hor_4 != None):
                        horizon4Paragraph.add_run(MottleShape_1_Text_English_Hor_4) # Mottle Shape
                        horizon4Paragraph.add_run(", ")
                    horizon4Paragraph.add_run(MottleColor_1_Text_English_Hor_4) # Mottle Color English
                    horizon4Paragraph.add_run(" (")
                    horizon4Paragraph.add_run(MottleColor_1_Code_Hor_4) # Mottle Color Code (Hue+Value+/+Chroma)
                    horizon4Paragraph.add_run(")")
                    horizon4Paragraph.add_run(" mottles; ")


                # Horizon Boundary
                if (BoundaryDistText_English_Hor_4 != None):    
                    horizon4Paragraph.add_run(BoundaryDistText_English_Hor_4) # horizon boundary dist
                    horizon4Paragraph.add_run(" ")
                    horizon4Paragraph.add_run(BoundaryTopoText_English_Hor_4) # horizon boundary topo
                    horizon4Paragraph.add_run(" boundary.")

                
                
                ### HORIZON 5 implementation ###

                # setting another paragraph because it have different format from above
                horizon5Paragraph = document.add_paragraph()
                # apply the font of Tahoma, size= 11, leftIndent = 2.45, first_line_indent = - 5 from left_indent,
                # align with justify_low, spaceBefore and spaceAfter = Pt(0)
                horizon5Paragraph.style = document.styles['Normal']
                horizon5ParagraphFormat = horizon5Paragraph.paragraph_format
                horizon5ParagraphFormat.line_spacing = 1
                horizon5ParagraphFormat.space_before = Pt(6)
                horizon5ParagraphFormat.space_after = Pt(0)
                horizon5Paragraph.paragraph_format.left_indent = Inches (2.5)
                horizon5Paragraph.paragraph_format.first_line_indent = Inches(-0.5)
                horizon5ParagraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Horison 
                if (horizonRow != -1):
                    if (HorDesignMasterText_English_Hor_5 != ""):
                        horizon5Paragraph.add_run(HorDesignDisconText_English_Hor_5 + HorDesignMasterText_English_Hor_5 + HorDesignSubText_English_Hor_5 + HorDesignNumberText_English_Hor_5)# horison design desig+master+sub+number
                        horizon5Paragraph.add_run("\t")
                    
                    # Depth
                    if (HorUpperToText_English_Hor_5 != ""):
                        horizon5Paragraph.add_run(HorUpperToText_English_Hor_5)# Upper limit to
                        horizon5Paragraph.add_run("-")
                    if (HorLowerToText_English_Hor_5 != ""):
                        horizon5Paragraph.add_run(HorLowerToText_English_Hor_5)# Lower limit to
                        horizon5Paragraph.add_run(" cm")
                        horizon5Paragraph.add_run("; ")

                # Soil Color 1
                if (SoilColor_1_Code_Hor_5 != "/" and SoilColor_1_Text_English_Hor_5 != None):
                    horizon5Paragraph.add_run(SoilColor_1_Text_English_Hor_5) # munsel color text english 
                    horizon5Paragraph.add_run(" (") 
                    horizon5Paragraph.add_run(SoilColor_1_Code_Hor_5) # Munsel Hue + Value + / + Chroma
                    horizon5Paragraph.add_run("); ") # 

                # Soil Color 2
                if (SoilColor_2_Code_Hor_5 != "/" and SoilColor_2_Text_English_Hor_5 != None):
                    horizon5Paragraph.add_run(SoilColor_2_Text_English_Hor_5) # munsel color text english
                    horizon5Paragraph.add_run(" (") 
                    horizon5Paragraph.add_run(SoilColor_2_Code_Hor_5) # Munsel Hue + Value + / + Chroma
                    horizon5Paragraph.add_run("); ") # 

                # Soil Color 3
                if (SoilColor_3_Code_Hor_5 != "/" and SoilColor_3_Text_English_Hor_5 != None):
                    horizon5Paragraph.add_run(SoilColor_3_Text_English_Hor_5) # munsel color text english
                    horizon5Paragraph.add_run(" (") 
                    horizon5Paragraph.add_run(SoilColor_3_Code_Hor_5) # Munsel Hue + Value + / + Chroma
                    horizon5Paragraph.add_run("); ") # 

                # Texture
                if (TextureClassText_English_Hor_5 != None):
                    horizon5Paragraph.add_run(TextureClassText_English_Hor_5) # texture class
                    horizon5Paragraph.add_run("; ")

                # Texture Sand
                if (TextureSandText_English_Hor_5 != None):
                    horizon5Paragraph.add_run(TextureSandText_English_Hor_5) # Texture sand
                    horizon5Paragraph.add_run(" sand; ")

                
                # Texture Modifier Size
                if(TextureModSizeText_English_Hor_5 != None):
                    horizon5Paragraph.add_run(TextureModSizeText_English_Hor_5)# Texture modifier size
                    horizon5Paragraph.add_run(", ")
                
                # Texture Modifier Abundance
                if(TextureModAbundanceText_English_Hor_5 != None):
                    horizon5Paragraph.add_run(TextureModAbundanceText_English_Hor_5)# Texture modifier size
                    horizon5Paragraph.add_run(" modifier; ")


                # Structure 1
                if (StructureShape_1_Text_English_Hor_5 != None):
                    horizon5Paragraph.add_run(StructureGrade_1_Text_English_Hor_5) # structure grade
                    horizon5Paragraph.add_run(", ")
                    horizon5Paragraph.add_run(StructureSize_1_Text_English_Hor_5) # structure size
                    horizon5Paragraph.add_run(", ")
                    horizon5Paragraph.add_run(StructureShape_1_Text_English_Hor_5) # structure shape
                    horizon5Paragraph.add_run(" structure; ")

                # Structure Relation
                if (StructureRelationText_English_Hor_5 != None):
                    horizon5Paragraph.add_run(StructureRelationText_English_Hor_5) # Structure Relation   

                # Structure 2
                if (StructureShape_2_Text_English_Hor_5 != None):
                    horizon5Paragraph.add_run(StructureGrade_2_Text_English_Hor_5) # structure grade
                    horizon5Paragraph.add_run(", ")
                    horizon5Paragraph.add_run(StructureSize_2_Text_English_Hor_5) # structure size
                    horizon5Paragraph.add_run(", ")
                    horizon5Paragraph.add_run(StructureShape_2_Text_English_Hor_5) # structure shape
                    horizon5Paragraph.add_run(" structure; ")

                # Consistence
                if (ConsistenceMoistText_English_Hor_5 != None):
                    horizon5Paragraph.add_run(ConsistenceMoistText_English_Hor_5) # consistency moist
                    horizon5Paragraph.add_run(" (moist), ")
                    horizon5Paragraph.add_run(ConsistenceStickinessText_English_Hor_5) # consistency stickiness
                    horizon5Paragraph.add_run(" and ")
                    horizon5Paragraph.add_run(ConsistencePlasticityText_English_Hor_5) # consitency plasticity
                    horizon5Paragraph.add_run(" (wet) consistency; ")
                
                # Roots
                if (RootFineText_English_Hor_5 != None):
                    horizon5Paragraph.add_run(RootFineText_English_Hor_5) # root fine
                    horizon5Paragraph.add_run(" fine")
                    if (RootMediumText_English_Hor_5 != None):
                        horizon5Paragraph.add_run(", ")
                        horizon5Paragraph.add_run(RootMediumText_English_Hor_5) # root medium
                        horizon5Paragraph.add_run(" medium")
                    if (RootCoarseText_English_Hor_5 != None):
                        horizon5Paragraph.add_run(", ")
                        horizon5Paragraph.add_run(RootCoarseText_English_Hor_5) # root coarse
                        horizon5Paragraph.add_run(" coarse")

                    horizon5Paragraph.add_run(" roots; ")

                # Mottles
                if (MottleColor_1_Text_English_Hor_5 != None):
                    if (MottleAbundance_1_Text_English_Hor_5 != None):
                        horizon5Paragraph.add_run(MottleAbundance_1_Text_English_Hor_5) # Mottle Abundance
                        horizon5Paragraph.add_run(", ")
                    if (MottleSize_1_Text_English_Hor_5 != None):
                        horizon5Paragraph.add_run(MottleSize_1_Text_English_Hor_5) # Mottle Size
                        horizon5Paragraph.add_run(", ")
                    if (MottleContrast_1_Text_English_Hor_5 != None):
                        horizon5Paragraph.add_run(MottleContrast_1_Text_English_Hor_5) # Mottle Contrast
                        horizon5Paragraph.add_run(", ")
                    if (MottleShape_1_Text_English_Hor_5 != None):
                        horizon5Paragraph.add_run(MottleShape_1_Text_English_Hor_5) # Mottle Shape
                        horizon5Paragraph.add_run(", ")
                    horizon5Paragraph.add_run(MottleColor_1_Text_English_Hor_5) # Mottle Color English
                    horizon5Paragraph.add_run(" (")
                    horizon5Paragraph.add_run(MottleColor_1_Code_Hor_5) # Mottle Color Code (Hue+Value+/+Chroma)
                    horizon5Paragraph.add_run(")")
                    horizon5Paragraph.add_run(" mottles; ")


                # Horizon Boundary
                if (BoundaryDistText_English_Hor_5 != None):    
                    horizon5Paragraph.add_run(BoundaryDistText_English_Hor_5) # horizon boundary dist
                    horizon5Paragraph.add_run(" ")
                    horizon5Paragraph.add_run(BoundaryTopoText_English_Hor_5) # horizon boundary topo
                    horizon5Paragraph.add_run(" boundary.")

                
                
                ### HORIZON 6 implementation ###

                # setting another paragraph because it have different format from above
                horizon6Paragraph = document.add_paragraph()
                # apply the font of Tahoma, size= 11, leftIndent = 2.45, first_line_indent = - 5 from left_indent,
                # align with justify_low, spaceBefore and spaceAfter = Pt(0)
                horizon6Paragraph.style = document.styles['Normal']
                horizon6ParagraphFormat = horizon6Paragraph.paragraph_format
                horizon6ParagraphFormat.line_spacing = 1
                horizon6ParagraphFormat.space_before = Pt(6)
                horizon6ParagraphFormat.space_after = Pt(0)
                horizon6Paragraph.paragraph_format.left_indent = Inches (2.5)
                horizon6Paragraph.paragraph_format.first_line_indent = Inches(-0.5)
                horizon6ParagraphFormat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Horizon 
                if (horizonRow != -1):
                    print("")
                    # if (HorDesignMasterText_English_Hor_6 != "" or HorDesignMasterText_English_Hor_6 != None):
                    horizon6Paragraph.add_run(HorDesignDisconText_English_Hor_6 + HorDesignMasterText_English_Hor_6 + HorDesignSubText_English_Hor_6 + HorDesignNumberText_English_Hor_6)# horison design desig+master+sub+number
                    horizon6Paragraph.add_run("\t")
                    # Depth
                    if (HorUpperToText_English_Hor_6 != ""):
                        horizon6Paragraph.add_run(HorUpperToText_English_Hor_6)# Upper limit to
                        horizon6Paragraph.add_run("-")
                    if (HorLowerToText_English_Hor_6 != ""):
                        horizon6Paragraph.add_run(HorLowerToText_English_Hor_6)# Lower limit to
                        horizon6Paragraph.add_run(" cm")
                        horizon6Paragraph.add_run("; ")

                # Soil Color 1
                if (SoilColor_1_Code_Hor_6 != "/" and SoilColor_1_Text_English_Hor_6 != None):
                    horizon6Paragraph.add_run(SoilColor_1_Text_English_Hor_6) # munsel color text english 
                    horizon6Paragraph.add_run(" (") 
                    horizon6Paragraph.add_run(SoilColor_1_Code_Hor_6) # Munsel Hue + Value + / + Chroma
                    horizon6Paragraph.add_run("); ") # 

                # Soil Color 2
                if (SoilColor_2_Code_Hor_6 != "/" and SoilColor_2_Text_English_Hor_6 != None):
                    horizon6Paragraph.add_run(SoilColor_2_Text_English_Hor_6) # munsel color text english
                    horizon6Paragraph.add_run(" (") 
                    horizon6Paragraph.add_run(SoilColor_2_Code_Hor_6) # Munsel Hue + Value + / + Chroma
                    horizon6Paragraph.add_run("); ") # 

                # Soil Color 3
                if (SoilColor_3_Code_Hor_6 != "/" and SoilColor_3_Text_English_Hor_6 != None):
                    horizon6Paragraph.add_run(SoilColor_3_Text_English_Hor_6) # munsel color text english
                    horizon6Paragraph.add_run(" (") 
                    horizon6Paragraph.add_run(SoilColor_3_Code_Hor_6) # Munsel Hue + Value + / + Chroma
                    horizon6Paragraph.add_run("); ") # 

                # Texture
                if (TextureClassText_English_Hor_6 != None):
                    horizon6Paragraph.add_run(TextureClassText_English_Hor_6) # texture class
                    horizon6Paragraph.add_run("; ")

                # Texture Sand
                if (TextureSandText_English_Hor_6 != None):
                    horizon6Paragraph.add_run(TextureSandText_English_Hor_6) # Texture sand
                    horizon6Paragraph.add_run(" sand; ")

                
                # Texture Modifier Size
                if(TextureModSizeText_English_Hor_6 != None):
                    horizon6Paragraph.add_run(TextureModSizeText_English_Hor_6)# Texture modifier size
                    horizon6Paragraph.add_run(", ")
                
                # Texture Modifier Abundance
                if(TextureModAbundanceText_English_Hor_6 != None):
                    horizon6Paragraph.add_run(TextureModAbundanceText_English_Hor_6)# Texture modifier size
                    horizon6Paragraph.add_run(" modifier; ")

                # Structure 1
                if (StructureShape_1_Text_English_Hor_6 != None):
                    horizon6Paragraph.add_run(StructureGrade_1_Text_English_Hor_6) # structure grade
                    horizon6Paragraph.add_run(", ")
                    horizon6Paragraph.add_run(StructureSize_1_Text_English_Hor_6) # structure size
                    horizon6Paragraph.add_run(", ")
                    horizon6Paragraph.add_run(StructureShape_1_Text_English_Hor_6) # structure shape
                    horizon6Paragraph.add_run(" structure; ")

                # Structure Relation
                if (StructureRelationText_English_Hor_6 != None):
                    horizon6Paragraph.add_run(StructureRelationText_English_Hor_6) # Structure Relation   

                # Structure 2
                if (StructureShape_2_Text_English_Hor_6 != None):
                    horizon6Paragraph.add_run(StructureGrade_2_Text_English_Hor_6) # structure grade
                    horizon6Paragraph.add_run(", ")
                    horizon6Paragraph.add_run(StructureSize_2_Text_English_Hor_6) # structure size
                    horizon6Paragraph.add_run(", ")
                    horizon6Paragraph.add_run(StructureShape_2_Text_English_Hor_6) # structure shape
                    horizon6Paragraph.add_run(" structure; ")

                # Consistence
                if (ConsistenceMoistText_English_Hor_6 != None):
                    horizon6Paragraph.add_run(ConsistenceMoistText_English_Hor_6) # consistency moist
                    horizon6Paragraph.add_run(" (moist), ")
                    horizon6Paragraph.add_run(ConsistenceStickinessText_English_Hor_6) # consistency stickiness
                    horizon6Paragraph.add_run(" and ")
                    horizon6Paragraph.add_run(ConsistencePlasticityText_English_Hor_6) # consitency plasticity
                    horizon6Paragraph.add_run(" (wet) consistency; ")
                
                # Roots
                if (RootFineText_English_Hor_6 != None):
                    horizon6Paragraph.add_run(RootFineText_English_Hor_6) # root fine
                    horizon6Paragraph.add_run(" fine")
                    if (RootMediumText_English_Hor_6 != None):
                        horizon6Paragraph.add_run(", ")
                        horizon6Paragraph.add_run(RootMediumText_English_Hor_6) # root medium
                        horizon6Paragraph.add_run(" medium")
                    if (RootCoarseText_English_Hor_6 != None):
                        horizon6Paragraph.add_run(", ")
                        horizon6Paragraph.add_run(RootCoarseText_English_Hor_6) # root coarse
                        horizon6Paragraph.add_run(" coarse")

                    horizon6Paragraph.add_run(" roots; ")

                # Mottles
                if (MottleColor_1_Text_English_Hor_6 != None):
                    if (MottleAbundance_1_Text_English_Hor_6 != None):
                        horizon6Paragraph.add_run(MottleAbundance_1_Text_English_Hor_6) # Mottle Abundance
                        horizon6Paragraph.add_run(", ")
                    if (MottleSize_1_Text_English_Hor_6 != None):
                        horizon6Paragraph.add_run(MottleSize_1_Text_English_Hor_6) # Mottle Size
                        horizon6Paragraph.add_run(", ")
                    if (MottleContrast_1_Text_English_Hor_6 != None):
                        horizon6Paragraph.add_run(MottleContrast_1_Text_English_Hor_6) # Mottle Contrast
                        horizon6Paragraph.add_run(", ")
                    if (MottleShape_1_Text_English_Hor_6 != None):
                        horizon6Paragraph.add_run(MottleShape_1_Text_English_Hor_6) # Mottle Shape
                        horizon6Paragraph.add_run(", ")
                    horizon6Paragraph.add_run(MottleColor_1_Text_English_Hor_6) # Mottle Color English
                    horizon6Paragraph.add_run(" (")
                    horizon6Paragraph.add_run(MottleColor_1_Code_Hor_6) # Mottle Color Code (Hue+Value+/+Chroma)
                    horizon6Paragraph.add_run(")")
                    horizon6Paragraph.add_run(" mottles; ")


                # Horizon Boundary
                if (BoundaryDistText_English_Hor_6 != None):    
                    horizon6Paragraph.add_run(BoundaryDistText_English_Hor_6) # horizon boundary dist
                    horizon6Paragraph.add_run(" ")
                    horizon6Paragraph.add_run(BoundaryTopoText_English_Hor_6) # horizon boundary topo
                    horizon6Paragraph.add_run(" boundary.")


                ### Add Picture method
                # document.add_picture("picture/FM066.jpg", width=Inches(2.5))
                
                document.add_page_break()

        else:
            # message show when the range is empty
            message = QMessageBox()
            message.setText("Please enter the range of number form!")
            message.setWindowTitle("Error!")
            message.setStandardButtons(QMessageBox.Ok)
            message.exec_()

        # save the dialog based on desired name with docx format
        document.save(self.saveDocxFile.text())
        # close the dialog after clicked the ok button
        self.exportDialog.close()

    def setSaveDocx(self):
        logging.info("set save docx")

        ### dialog to get save file name
        options = QFileDialog.Options()
        
        fileName, _ = QFileDialog.getSaveFileName(self,
                "QFileDialog.getSaveFileName()",
                self.saveDocxFile.text(),
                "Docx Files (*.docx)", options=options)
        if fileName:
            self.saveDocxFile.setText(fileName)
    
    def setSaveLocation(self):
        logging.info("set save location")

        ### dialog to get save file name for SQlite Database
        options = QFileDialog.Options()
        
        fileName, _ = QFileDialog.getSaveFileName(self,
                "QFileDialog.getSaveFileName()",
                self.saveSqliteDBLineEdit.text(),
                "Database Files (*.db)", options=options)
        if fileName:
            self.saveSqliteDBLineEdit.setText(fileName)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    
    # connect to the database, if connection is invalid/ cant connect close the window
    if not connection.createConnection():
        sys.exit(1)

    main = First()
    # main.setWindowState(Qt.WindowMaximized)
    main.show()

    sys.exit(app.exec_())
